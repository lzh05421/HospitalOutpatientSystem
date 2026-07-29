from __future__ import annotations

import argparse
from pathlib import Path
import json
import math

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSET_DIR = DOCS / "course_design_assets"
TEMPLATE = DOCS / "course_design_template.docx"
OUTPUT = DOCS / "医院门诊挂号与药品管理系统-课程设计总结.docx"
APPENDIX_OUTPUT = DOCS / "医院门诊挂号与药品管理系统-课程设计总结-附件.docx"
REFERENCES_OUTPUT = DOCS / "医院门诊挂号与药品管理系统-课程设计总结-文献核验清单.json"
SUMMARY_MD = DOCS / "医院门诊挂号与药品管理系统-课程设计总结.md"
ACTIVE_TEMPLATE = TEMPLATE
ACTIVE_OUTPUT = OUTPUT
ACTIVE_APPENDIX_OUTPUT = APPENDIX_OUTPUT
ACTIVE_REFERENCES_OUTPUT = REFERENCES_OUTPUT
ACTIVE_SUMMARY_MD = SUMMARY_MD

TITLE = "基于Qt和MySQL的医院门诊挂号与药品管理系统"
DATE_TEXT = "2026年6月29日"
STUDENT_NAME = "刘子航"
ADVISOR_NAME = "________"


def font_path(*names: str) -> str:
    candidates = [
        Path("C:/Windows/Fonts") / name
        for name in names
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return str(Path("C:/Windows/Fonts/msyh.ttc"))


FONT_REG = font_path("simsun.ttc", "msyh.ttc")
FONT_BOLD = font_path("simhei.ttf", "msyhbd.ttc", "msyh.ttc")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def draw_text_center(draw: ImageDraw.ImageDraw, xy, text: str, font, fill="#1f2937"):
    x, y = xy
    bbox = draw.textbbox((0, 0), text, font=font)
    draw.text((x - (bbox[2] - bbox[0]) / 2, y - (bbox[3] - bbox[1]) / 2), text, font=font, fill=fill)


def rounded_rect(draw, box, fill, outline="#334155", radius=14, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#334155", width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    points = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6)),
        (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(points, fill=fill)


def save_architecture_diagram(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(28, True)
    body = pil_font(23)
    small = pil_font(20)
    draw_text_center(draw, (800, 55), "系统总体架构图", title, "#0f172a")

    blocks = [
        ((80, 165, 390, 395), "#dbeafe", "医院工作站客户端", ["Qt Widgets界面", "挂号员/医生/药师/财务/管理员", "模块导航与表格操作"]),
        ((80, 505, 390, 735), "#e0f2fe", "患者端入口", ["患者注册登录", "预约挂号", "个人信息维护"]),
        ((535, 210, 820, 690), "#dcfce7", "TCP通信层", ["JSON Lines协议", "请求/响应封装", "会话Token传递"]),
        ((960, 145, 1320, 755), "#ede9fe", "Qt服务端", ["RequestRouter路由", "SessionManager会话", "Authorization权限校验", "ModuleService业务处理", "OperationLog审计"]),
        ((1395, 180, 1540, 360), "#fee2e2", "MySQL", ["业务数据", "关系约束"]),
        ((1395, 430, 1540, 610), "#fef3c7", "Redis", ["可选缓存", "库存扣减脚本"]),
    ]
    for box, fill, name, lines in blocks:
        rounded_rect(draw, box, fill, "#475569", 18, 3)
        draw_text_center(draw, ((box[0] + box[2]) / 2, box[1] + 42), name, h, "#111827")
        for i, line in enumerate(lines):
            draw_text_center(draw, ((box[0] + box[2]) / 2, box[1] + 95 + i * 42), line, body if i < 2 else small, "#1f2937")

    arrow(draw, (390, 280), (535, 340))
    arrow(draw, (390, 620), (535, 560))
    arrow(draw, (820, 450), (960, 450))
    arrow(draw, (1320, 300), (1395, 275))
    arrow(draw, (1320, 520), (1395, 520))
    draw_text_center(draw, (462, 260), "请求", small, "#475569")
    draw_text_center(draw, (890, 415), "module/action/payload", small, "#475569")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def save_function_diagram(path: Path):
    img = Image.new("RGB", (1700, 1100), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(25, True)
    body = pil_font(21)
    draw_text_center(draw, (850, 55), "系统功能架构图", title, "#0f172a")
    rounded_rect(draw, (600, 105, 1100, 180), "#0f766e", "#0f766e", 18, 2)
    draw_text_center(draw, (850, 143), "医院门诊挂号与药品管理系统", h, "#ffffff")

    groups = [
        ("基础管理", ["患者管理", "科室管理", "医生管理", "权限配置"], "#dbeafe"),
        ("挂号排班", ["医生排班", "预约挂号", "候诊队列", "退号处理"], "#dcfce7"),
        ("诊疗业务", ["医生接诊", "病历档案", "检查检验", "处方开立"], "#ede9fe"),
        ("药品收费", ["处方审核", "药品库存", "收费结算", "退费审核"], "#fef3c7"),
        ("统计审计", ["院长驾驶舱", "费用统计", "操作日志", "库存预警"], "#fee2e2"),
    ]
    x_positions = [80, 405, 730, 1055, 1380]
    for x, (name, items, fill) in zip(x_positions, groups):
        arrow(draw, (850, 180), (x + 120, 265), "#64748b", 3)
        rounded_rect(draw, (x, 265, x + 240, 345), fill, "#334155", 14, 2)
        draw_text_center(draw, (x + 120, 305), name, h, "#111827")
        y = 390
        for item in items:
            rounded_rect(draw, (x, y, x + 240, y + 62), "#f8fafc", "#94a3b8", 10, 2)
            draw_text_center(draw, (x + 120, y + 31), item, body, "#1f2937")
            y += 82
    img.save(path)


def save_tech_diagram(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(26, True)
    body = pil_font(22)
    draw_text_center(draw, (800, 55), "技术架构图", title, "#0f172a")
    layers = [
        ("表现层", "Qt Widgets、QMainWindow、QStackedWidget、ModulePage、表格与表单控件", "#dbeafe"),
        ("通信层", "QTcpSocket、JSON Lines协议、Request/Response封装、同步与异步回调", "#dcfce7"),
        ("服务层", "HospitalServer、ClientConnection、RequestRouter、ModuleService、权限校验", "#ede9fe"),
        ("业务层", "挂号、排班、接诊、检查、处方、库存、收费、统计、权限、日志", "#fef3c7"),
        ("数据层", "MySQL关系数据库、事务处理、外键约束、可选Redis库存脚本", "#fee2e2"),
        ("工程支撑", "CMake、Qt5/Qt6兼容、Windows/Linux/银河麒麟脚本、单元与源码测试", "#e2e8f0"),
    ]
    y = 125
    for name, desc, fill in layers:
        rounded_rect(draw, (150, y, 1450, y + 95), fill, "#475569", 14, 3)
        draw_text_center(draw, (285, y + 48), name, h, "#111827")
        draw.text((420, y + 33), desc, font=body, fill="#1f2937")
        if y < 710:
            arrow(draw, (800, y + 95), (800, y + 125), "#64748b", 3)
        y += 125
    img.save(path)


def save_initial_class_diagram(path: Path):
    img = Image.new("RGB", (1700, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(24, True)
    body = pil_font(20)
    draw_text_center(draw, (850, 55), "初步类图设计", title, "#0f172a")
    classes = [
        ((80, 150, 390, 350), "MainWindow", ["m_apiClient", "m_navigation", "addModulePage()", "canAccess()"], "#dbeafe"),
        ((80, 500, 390, 740), "ModulePage", ["m_apiClient", "m_table", "refresh()", "rowsUpdated()"], "#dbeafe"),
        ((500, 150, 810, 350), "ApiClient", ["connectToServer()", "sendRequest()", "onReadyRead()"], "#dcfce7"),
        ((910, 150, 1220, 350), "HospitalServer", ["listen()", "newConnection()", "registerServices()"], "#ede9fe"),
        ((1300, 150, 1610, 390), "RequestRouter", ["registerService()", "route()", "authorize()", "isAllowed()"], "#ede9fe"),
        ((910, 520, 1220, 760), "ModuleService", ["handle(request)", "业务服务基类"], "#fef3c7"),
        ((1300, 520, 1610, 760), "DatabaseManager", ["open()", "query()", "transaction()", "reconnect()"], "#fee2e2"),
    ]
    for box, name, members, fill in classes:
        rounded_rect(draw, box, fill, "#475569", 10, 3)
        draw.line((box[0], box[1] + 52, box[2], box[1] + 52), fill="#475569", width=2)
        draw_text_center(draw, ((box[0] + box[2]) / 2, box[1] + 28), name, h, "#111827")
        for i, m in enumerate(members):
            draw.text((box[0] + 24, box[1] + 72 + i * 36), m, font=body, fill="#1f2937")
    arrow(draw, (390, 250), (500, 250))
    arrow(draw, (390, 620), (500, 290))
    arrow(draw, (810, 250), (910, 250))
    arrow(draw, (1220, 250), (1300, 250))
    arrow(draw, (1455, 390), (1455, 520))
    arrow(draw, (1065, 520), (1300, 640))
    draw_text_center(draw, (445, 225), "调用", body, "#475569")
    draw_text_center(draw, (860, 225), "TCP", body, "#475569")
    draw_text_center(draw, (1260, 225), "路由", body, "#475569")
    img.save(path)


def save_sequence_diagram(path: Path):
    img = Image.new("RGB", (1700, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(24, True)
    body = pil_font(19)
    draw_text_center(draw, (850, 55), "挂号业务时序图", title, "#0f172a")
    actors = [
        ("挂号员/患者", 160),
        ("RegistrationPage", 440),
        ("ApiClient", 710),
        ("RequestRouter", 980),
        ("RegistrationService", 1260),
        ("MySQL", 1530),
    ]
    for name, x in actors:
        rounded_rect(draw, (x - 105, 120, x + 105, 175), "#e2e8f0", "#475569", 10, 2)
        draw_text_center(draw, (x, 148), name, h, "#111827")
        draw.line((x, 175, x, 900), fill="#94a3b8", width=2)
    steps = [
        (160, 440, "选择患者、科室、医生和号源"),
        (440, 710, "sendRequest(registration/create)"),
        (710, 980, "JSON Lines请求"),
        (980, 1260, "鉴权后转发"),
        (1260, 1530, "校验号源、患者、医保状态"),
        (1530, 1260, "返回可用号源与事务结果"),
        (1260, 1530, "写入registrations并更新remain_quota"),
        (1530, 1260, "提交事务"),
        (1260, 980, "生成账单和业务响应"),
        (980, 710, "统一响应"),
        (710, 440, "onResponseReceived"),
        (440, 160, "刷新列表并提示挂号成功"),
    ]
    y = 230
    for i, (x1, x2, text) in enumerate(steps, 1):
        arrow(draw, (x1, y), (x2, y), "#334155", 3)
        label_y = y - 33 if i % 2 else y + 8
        draw_text_center(draw, ((x1 + x2) / 2, label_y), f"{i}. {text}", body, "#1f2937")
        y += 58
    img.save(path)


def save_detailed_class_diagram(path: Path):
    img = Image.new("RGB", (1700, 1050), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(23, True)
    body = pil_font(19)
    draw_text_center(draw, (850, 55), "详细类图设计", title, "#0f172a")
    services = [
        ("AuthService", "登录认证"),
        ("PatientService", "患者档案"),
        ("RegistrationService", "挂号退号"),
        ("ScheduleService", "医生排班"),
        ("ConsultationService", "候诊接诊"),
        ("PrescriptionService", "处方审核发药"),
        ("InventoryService", "库存出入库"),
        ("BillingService", "收费退费"),
        ("StatisticsService", "费用统计"),
        ("PermissionAdminService", "权限配置"),
    ]
    rounded_rect(draw, (630, 130, 1070, 240), "#ede9fe", "#475569", 14, 3)
    draw_text_center(draw, (850, 165), "RequestRouter", h, "#111827")
    draw_text_center(draw, (850, 205), "route() / authorize() / writeOperationLog()", body, "#1f2937")
    rounded_rect(draw, (650, 845, 1050, 945), "#fee2e2", "#475569", 14, 3)
    draw_text_center(draw, (850, 880), "DatabaseManager", h, "#111827")
    draw_text_center(draw, (850, 918), "事务、查询、重连、演示数据回退", body, "#1f2937")
    for i, (name, desc) in enumerate(services):
        col = i % 5
        row = i // 5
        x = 75 + col * 325
        y = 360 + row * 230
        rounded_rect(draw, (x, y, x + 260, y + 130), "#f8fafc", "#64748b", 12, 2)
        draw_text_center(draw, (x + 130, y + 38), name, h, "#111827")
        draw_text_center(draw, (x + 130, y + 82), desc, body, "#1f2937")
        arrow(draw, (850, 240), (x + 130, y), "#64748b", 2)
        arrow(draw, (x + 130, y + 130), (850, 845), "#94a3b8", 2)
    img.save(path)


def save_er_diagram(path: Path):
    img = Image.new("RGB", (1900, 1250), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(23, True)
    body = pil_font(18)
    draw_text_center(draw, (950, 55), "数据库实体关系图", title, "#0f172a")
    entities = [
        ("users", "员工用户", 110, 150, "#dbeafe"),
        ("roles", "角色", 430, 150, "#dbeafe"),
        ("departments", "科室", 750, 150, "#dcfce7"),
        ("doctors", "医生", 1070, 150, "#dcfce7"),
        ("patients", "患者", 1390, 150, "#fef3c7"),
        ("doctor_schedules", "排班号源", 750, 390, "#dcfce7"),
        ("registrations", "挂号记录", 1070, 390, "#fef3c7"),
        ("medical_records", "电子病历", 1390, 390, "#ede9fe"),
        ("examinations", "检查检验", 430, 660, "#ede9fe"),
        ("prescriptions", "处方", 750, 660, "#ede9fe"),
        ("prescription_items", "处方明细", 1070, 660, "#ede9fe"),
        ("drugs", "药品", 1390, 660, "#fee2e2"),
        ("stock_records", "库存流水", 1390, 910, "#fee2e2"),
        ("bills", "账单", 750, 910, "#fef3c7"),
        ("payments", "支付记录", 1070, 910, "#fef3c7"),
        ("operation_logs", "操作日志", 110, 910, "#e2e8f0"),
    ]
    centers = {}
    for table, name, x, y, fill in entities:
        rounded_rect(draw, (x, y, x + 260, y + 115), fill, "#475569", 10, 2)
        draw_text_center(draw, (x + 130, y + 36), table, h, "#111827")
        draw_text_center(draw, (x + 130, y + 78), name, body, "#1f2937")
        centers[table] = (x + 130, y + 58)
    links = [
        ("roles", "users"), ("departments", "doctors"), ("users", "doctors"),
        ("doctors", "doctor_schedules"), ("departments", "doctor_schedules"),
        ("patients", "registrations"), ("doctors", "registrations"), ("doctor_schedules", "registrations"),
        ("registrations", "medical_records"), ("registrations", "examinations"),
        ("registrations", "prescriptions"), ("prescriptions", "prescription_items"),
        ("drugs", "prescription_items"), ("drugs", "stock_records"),
        ("registrations", "bills"), ("patients", "bills"), ("bills", "payments"),
        ("users", "operation_logs"),
    ]
    for a, b in links:
        arrow(draw, centers[a], centers[b], "#64748b", 3)
    draw.text((95, 1110), "说明：图中保留核心业务实体。权限、医保、审计明细、检查项目、药品分类等扩展表通过外键与核心实体关联。", font=body, fill="#334155")
    img.save(path)


def save_core_er_diagram(path: Path):
    img = Image.new("RGB", (1700, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(23, True)
    body = pil_font(18)
    draw_text_center(draw, (850, 55), "核心业务 E-R 图", title, "#0f172a")
    entities = [
        ("patients", "患者\npatient_id, patient_no, id_card, phone", 80, 140, "#fef3c7"),
        ("departments", "科室\ndept_id, dept_code, dept_name", 500, 140, "#dcfce7"),
        ("doctors", "医生\ndoctor_id, dept_id, title, fee", 920, 140, "#dcfce7"),
        ("doctor_schedules", "排班号源\nschedule_id, work_date, period, remain_quota", 500, 365, "#dcfce7"),
        ("registrations", "挂号\nregistration_id, registration_no, status, fee", 920, 365, "#fef3c7"),
        ("medical_records", "病历\nrecord_id, complaint, diagnosis, advice", 1280, 365, "#ede9fe"),
        ("prescriptions", "处方\nprescription_id, status, total_amount", 500, 620, "#ede9fe"),
        ("prescription_items", "处方明细\ndrug_id, quantity, amount", 920, 620, "#ede9fe"),
        ("drugs", "药品\ndrug_id, drug_code, stock_quantity", 1280, 620, "#fee2e2"),
        ("bills", "账单\nbill_id, bill_no, status, total_amount", 500, 820, "#fef3c7"),
        ("payments", "支付\npayment_id, pay_method, pay_status", 920, 820, "#fef3c7"),
    ]
    centers = {}
    for table, desc, x, y, fill in entities:
        rounded_rect(draw, (x, y, x + 330, y + 120), fill, "#475569", 10, 2)
        name, detail = desc.split("\n", 1)
        draw_text_center(draw, (x + 165, y + 34), table, h, "#111827")
        draw_text_center(draw, (x + 165, y + 66), name, body, "#1f2937")
        draw_text_center(draw, (x + 165, y + 94), detail, body, "#334155")
        centers[table] = (x + 165, y + 60)
    for a, b in [
        ("departments", "doctors"),
        ("doctors", "doctor_schedules"),
        ("patients", "registrations"),
        ("doctor_schedules", "registrations"),
        ("registrations", "medical_records"),
        ("registrations", "prescriptions"),
        ("prescriptions", "prescription_items"),
        ("drugs", "prescription_items"),
        ("registrations", "bills"),
        ("bills", "payments"),
    ]:
        arrow(draw, centers[a], centers[b], "#64748b", 3)
    img.save(path)


def save_permission_flow_diagram(path: Path):
    img = Image.new("RGB", (1600, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(24, True)
    body = pil_font(20)
    draw_text_center(draw, (800, 55), "登录权限控制流程图", title, "#0f172a")
    steps = [
        ("输入账号密码", "LoginDialog/PatientLoginDialog", 120, 180, "#dbeafe"),
        ("提交登录请求", "auth/login 或 patient/login", 430, 180, "#dcfce7"),
        ("校验用户状态", "AuthService + sys_user/users", 740, 180, "#ede9fe"),
        ("生成会话Token", "SessionManager保存角色与权限", 1050, 180, "#fef3c7"),
        ("进入主界面", "MainWindow按权限加载模块", 120, 520, "#dbeafe"),
        ("接口二次鉴权", "RequestRouter canAccess", 430, 520, "#ede9fe"),
        ("业务服务处理", "ModuleService handle", 740, 520, "#fef3c7"),
        ("写入操作日志", "operation_logs/audit_log_details", 1050, 520, "#fee2e2"),
    ]
    for i, (name, desc, x, y, fill) in enumerate(steps):
        rounded_rect(draw, (x, y, x + 250, y + 120), fill, "#475569", 12, 2)
        draw_text_center(draw, (x + 125, y + 38), name, h, "#111827")
        draw_text_center(draw, (x + 125, y + 82), desc, body, "#334155")
        if i < 3:
            arrow(draw, (x + 250, y + 60), (x + 310, y + 60), "#64748b", 3)
        elif i == 3:
            arrow(draw, (x + 125, y + 120), (245, 520), "#64748b", 3)
        elif i < 7:
            arrow(draw, (x + 250, y + 60), (x + 310, y + 60), "#64748b", 3)
    img.save(path)


def save_prescription_flow_diagram(path: Path):
    img = Image.new("RGB", (1650, 920), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(23, True)
    body = pil_font(19)
    draw_text_center(draw, (825, 55), "处方审核与发药流程图", title, "#0f172a")
    steps = [
        ("医生开立处方", "createPrescription", 80, 170, "#ede9fe"),
        ("规则审核", "pass_rules 检查", 390, 170, "#fef3c7"),
        ("药师复核", "reviewPrescription", 700, 170, "#dcfce7"),
        ("收费确认", "bill/payment 状态", 1010, 170, "#fef3c7"),
        ("发药扣库存", "dispensePrescription", 1320, 170, "#fee2e2"),
        ("处方退药", "returnPrescription", 700, 520, "#fee2e2"),
        ("库存流水", "stock_records", 1010, 520, "#fee2e2"),
        ("日志留痕", "operation_logs", 1320, 520, "#e2e8f0"),
    ]
    for i, (name, desc, x, y, fill) in enumerate(steps):
        rounded_rect(draw, (x, y, x + 240, y + 115), fill, "#475569", 12, 2)
        draw_text_center(draw, (x + 120, y + 36), name, h, "#111827")
        draw_text_center(draw, (x + 120, y + 78), desc, body, "#334155")
    for a, b in [(200, 510), (510, 820), (820, 1130), (1130, 1440)]:
        arrow(draw, (a, 228), (b - 120, 228), "#64748b", 3)
    arrow(draw, (1440, 285), (1440, 520), "#64748b", 3)
    arrow(draw, (1440, 578), (1250, 578), "#64748b", 3)
    arrow(draw, (1010, 578), (940, 578), "#64748b", 3)
    arrow(draw, (940, 578), (1320, 578), "#94a3b8", 2)
    img.save(path)


def save_payment_flow_diagram(path: Path):
    img = Image.new("RGB", (1650, 920), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(23, True)
    body = pil_font(19)
    draw_text_center(draw, (825, 55), "收费结算与退费流程图", title, "#0f172a")
    steps = [
        ("生成账单", "挂号/处方/检查形成费用", 100, 160, "#fef3c7"),
        ("选择支付方式", "自费/医保/扫码", 430, 160, "#dbeafe"),
        ("支付处理", "BillingService", 760, 160, "#ede9fe"),
        ("写入支付记录", "payments", 1090, 160, "#dcfce7"),
        ("更新账单状态", "bills.status=已支付", 1090, 430, "#fef3c7"),
        ("退费申请", "requestRefundBill", 430, 650, "#fee2e2"),
        ("退费审核", "reviewRefundBill", 760, 650, "#fee2e2"),
        ("审计记录", "audit_log_details", 1090, 650, "#e2e8f0"),
    ]
    for name, desc, x, y, fill in steps:
        rounded_rect(draw, (x, y, x + 250, y + 115), fill, "#475569", 12, 2)
        draw_text_center(draw, (x + 125, y + 36), name, h, "#111827")
        draw_text_center(draw, (x + 125, y + 78), desc, body, "#334155")
    for x in [350, 680, 1010]:
        arrow(draw, (x, 218), (x + 80, 218), "#64748b", 3)
    arrow(draw, (1215, 275), (1215, 430), "#64748b", 3)
    arrow(draw, (555, 650), (555, 520), "#94a3b8", 2)
    arrow(draw, (555, 650), (760, 708), "#64748b", 3)
    arrow(draw, (1010, 708), (1090, 708), "#64748b", 3)
    img.save(path)


def save_platform_ui_diagram(path: Path):
    img = Image.new("RGB", (1700, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(24, True)
    body = pil_font(19)
    draw_text_center(draw, (850, 55), "平台端主界面布局图", title, "#0f172a")
    rounded_rect(draw, (80, 120, 1620, 960), "#ffffff", "#475569", 16, 3)
    rounded_rect(draw, (80, 120, 350, 960), "#e0f2fe", "#475569", 16, 2)
    draw_text_center(draw, (215, 165), "模块导航", h, "#0f172a")
    navs = ["院长驾驶舱", "患者管理", "挂号管理", "候诊队列", "医生接诊", "处方管理", "药品库存", "收费结算", "权限配置"]
    y = 220
    for item in navs:
        rounded_rect(draw, (115, y, 315, y + 54), "#ffffff", "#94a3b8", 8, 2)
        draw_text_center(draw, (215, y + 27), item, body, "#1f2937")
        y += 72
    rounded_rect(draw, (390, 155, 1550, 235), "#f1f5f9", "#cbd5e1", 10, 2)
    draw_text_center(draw, (540, 195), "当前模块：挂号管理", h, "#111827")
    draw_text_center(draw, (1280, 195), "刷新  新增  修改  退号", body, "#334155")
    rounded_rect(draw, (390, 270, 1550, 420), "#f8fafc", "#cbd5e1", 10, 2)
    filters = ["患者姓名", "科室", "医生", "就诊日期", "状态"]
    x = 430
    for f in filters:
        rounded_rect(draw, (x, 310, x + 175, 365), "#ffffff", "#94a3b8", 8, 2)
        draw_text_center(draw, (x + 88, 338), f, body, "#334155")
        x += 210
    rounded_rect(draw, (390, 455, 1550, 900), "#ffffff", "#cbd5e1", 10, 2)
    headers = ["挂号单号", "患者", "科室", "医生", "时段", "状态", "费用"]
    x = 430
    for hd in headers:
        draw_text_center(draw, (x, 500), hd, body, "#111827")
        x += 155
    for row in range(5):
        y = 545 + row * 65
        draw.line((420, y, 1515, y), fill="#e2e8f0", width=2)
        x = 430
        for col in range(len(headers)):
            draw_text_center(draw, (x, y + 32), "示例数据", body, "#475569")
            x += 155
    img.save(path)


def save_patient_ui_diagram(path: Path):
    img = Image.new("RGB", (1700, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(24, True)
    body = pil_font(19)
    draw_text_center(draw, (850, 55), "患者端预约界面布局图", title, "#0f172a")
    rounded_rect(draw, (160, 130, 1540, 900), "#f8fafc", "#475569", 16, 3)
    rounded_rect(draw, (220, 190, 1480, 285), "#dbeafe", "#3b82f6", 14, 2)
    draw_text_center(draw, (430, 238), "患者预约挂号", h, "#0f172a")
    draw_text_center(draw, (1160, 238), "个人信息  我的挂号  退出登录", body, "#334155")
    labels = ["选择科室", "选择医生", "选择日期", "选择时段", "医保类型", "联系电话"]
    positions = [(260, 360), (650, 360), (1040, 360), (260, 500), (650, 500), (1040, 500)]
    for label, (x, y) in zip(labels, positions):
        draw_text_center(draw, (x + 90, y - 28), label, body, "#334155")
        rounded_rect(draw, (x, y, x + 280, y + 58), "#ffffff", "#94a3b8", 8, 2)
        draw_text_center(draw, (x + 140, y + 29), "请选择/请输入", body, "#64748b")
    rounded_rect(draw, (260, 645, 1440, 780), "#ffffff", "#cbd5e1", 10, 2)
    draw_text_center(draw, (430, 690), "可用号源：上午 12 个，下午 9 个", body, "#1f2937")
    draw_text_center(draw, (430, 735), "挂号费用：15.00 元", body, "#1f2937")
    rounded_rect(draw, (1110, 680, 1370, 750), "#0f766e", "#0f766e", 10, 2)
    draw_text_center(draw, (1240, 715), "提交预约", h, "#ffffff")
    img.save(path)


def save_directory_diagram(path: Path):
    img = Image.new("RGB", (1600, 900), "#ffffff")
    draw = ImageDraw.Draw(img)
    title = pil_font(42, True)
    h = pil_font(26, True)
    body = pil_font(22)
    draw_text_center(draw, (800, 55), "程序目录结构图", title, "#0f172a")
    root_box = (590, 115, 1010, 185)
    rounded_rect(draw, root_box, "#0f766e", "#0f766e", 14, 2)
    draw_text_center(draw, (800, 150), "HospitalOutpatientSystem", h, "#ffffff")
    items = [
        ("client", "Qt桌面客户端：页面、登录、患者端入口", 110, 310, "#dbeafe"),
        ("server", "Qt TCP服务端：路由、业务服务、数据库访问", 500, 310, "#ede9fe"),
        ("common", "通信协议：Request、Response、JSON封装", 890, 310, "#dcfce7"),
        ("database", "MySQL建库脚本、业务表和演示数据", 1280, 310, "#fee2e2"),
        ("config", "服务端连接配置与跨平台示例配置", 300, 610, "#fef3c7"),
        ("scripts", "Windows/Linux/银河麒麟构建运行脚本", 690, 610, "#e2e8f0"),
        ("tests", "协议、权限、排班、收费、库存等测试", 1080, 610, "#e0f2fe"),
    ]
    for name, desc, x, y, fill in items:
        arrow(draw, (800, 185), (x + 130, y), "#64748b", 3)
        rounded_rect(draw, (x, y, x + 260, y + 125), fill, "#475569", 12, 2)
        draw_text_center(draw, (x + 130, y + 36), name, h, "#111827")
        draw_text_center(draw, (x + 130, y + 82), desc[:18], body, "#1f2937")
        draw_text_center(draw, (x + 130, y + 108), desc[18:], body, "#1f2937")
    img.save(path)


def generate_images() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        "arch": ASSET_DIR / "figure_3_1_architecture.png",
        "function": ASSET_DIR / "figure_3_2_function.png",
        "tech": ASSET_DIR / "figure_3_3_tech.png",
        "class_initial": ASSET_DIR / "figure_3_4_initial_class.png",
        "sequence": ASSET_DIR / "figure_3_5_sequence.png",
        "class_detail": ASSET_DIR / "figure_3_6_detail_class.png",
        "er": ASSET_DIR / "figure_3_7_er.png",
        "core_er": ASSET_DIR / "figure_3_8_core_er.png",
        "directory": ASSET_DIR / "figure_4_1_directory.png",
        "permission_flow": ASSET_DIR / "figure_4_2_permission_flow.png",
        "prescription_flow": ASSET_DIR / "figure_4_3_prescription_flow.png",
        "payment_flow": ASSET_DIR / "figure_4_4_payment_flow.png",
        "platform_ui": ASSET_DIR / "figure_4_5_platform_ui.png",
        "patient_ui": ASSET_DIR / "figure_4_6_patient_ui.png",
    }
    save_architecture_diagram(images["arch"])
    save_function_diagram(images["function"])
    save_tech_diagram(images["tech"])
    save_initial_class_diagram(images["class_initial"])
    save_sequence_diagram(images["sequence"])
    save_detailed_class_diagram(images["class_detail"])
    save_er_diagram(images["er"])
    save_core_er_diagram(images["core_er"])
    save_directory_diagram(images["directory"])
    save_permission_flow_diagram(images["permission_flow"])
    save_prescription_flow_diagram(images["prescription_flow"])
    save_payment_flow_diagram(images["payment_flow"])
    save_platform_ui_diagram(images["platform_ui"])
    save_patient_ui_diagram(images["patient_ui"])
    return images


def set_run_font(run, size=12, bold=False, font="宋体", color: str | None = None):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, first_line=True, line=1.5, before=0, after=0, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(0.74) if first_line else None
    if align is not None:
        paragraph.alignment = align


def clear_body(doc: Document):
    body = doc._element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def configure_document(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    for style_name, size, font, bold in [
        ("Heading 1", 18, "黑体", True),
        ("Heading 2", 15, "黑体", True),
        ("Heading 3", 14, "黑体", True),
    ]:
        style = styles[style_name]
        style.font.name = font
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(12)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_paragraph(doc: Document, text: str = "", first_line=True, size=12, bold=False, font="宋体", align=None):
    p = doc.add_paragraph()
    set_para(p, first_line=first_line, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, font=font)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if level <= 3 else "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, 18, True, "黑体")
    elif level == 2:
        set_run_font(r, 15, True, "黑体")
    else:
        set_run_font(r, 14, True, "黑体")
    return p


def add_center(doc: Document, text: str, size=12, bold=False, font="宋体", after=0):
    p = doc.add_paragraph()
    set_para(p, first_line=False, after=after, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    set_run_font(r, size, bold, font)
    return p


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_para(p, first_line=False, line=1.5, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, font="宋体")


def set_table_borders(table, three_line=False):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if three_line and edge in ("left", "right", "insideV"):
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "12" if edge in ("top", "bottom") and three_line else "6")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def add_table_caption(doc: Document, caption: str):
    p = doc.add_paragraph()
    set_para(p, first_line=False, before=6, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(caption)
    set_run_font(r, 10.5, True, "黑体")
    return p


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], three_line=True):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=10.5)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5)
    set_table_borders(table, three_line=three_line)
    return table


def add_code_listing(doc: Document, caption: str, code: str):
    p = doc.add_paragraph()
    set_para(p, first_line=False, before=6, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(caption)
    set_run_font(r, 10.5, True, "黑体")
    for line in code.strip("\n").splitlines():
        para = doc.add_paragraph()
        set_para(para, first_line=False, line=1.0, before=0, after=0)
        run = para.add_run(line)
        set_run_font(run, 9, False, "Consolas")


def add_figure(doc: Document, image_path: Path, caption: str, width_cm=14.8):
    p = doc.add_paragraph()
    set_para(p, first_line=False, before=6, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    set_para(cap, first_line=False, before=3, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = cap.add_run(caption)
    set_run_font(r, 10.5, True, "黑体")


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document):
    p = doc.add_paragraph()
    set_para(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("目  录")
    set_run_font(r, 18, True, "黑体")
    p = doc.add_paragraph()
    set_para(p, first_line=False)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    text = OxmlElement("w:t")
    text.text = "请在 Word 中右键目录选择“更新域”以生成页码。"
    run._r.append(text)
    run._r.append(fld_end)
    add_page_break(doc)


def build_cover(doc: Document):
    add_center(doc, "哈尔滨信息工程学院", 22, True, "宋体", after=18)
    add_center(doc, "项目综合实践II课设总结", 22, True, "宋体", after=38)
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["题   目：", "学    院：", "年    级：", "专    业：", "姓    名：", "指导教师："]
    values = [TITLE, "软件学院", "2023级", "计算机科学与技术", STUDENT_NAME, ADVISOR_NAME]
    for row, label, value in zip(table.rows, labels, values):
        row.height = Cm(1.05)
        set_cell_text(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=14)
        set_cell_text(row.cells[1], value, align=WD_ALIGN_PARAGRAPH.LEFT, size=14)
    set_table_borders(table, three_line=False)
    for _ in range(5):
        add_paragraph(doc, "", first_line=False)
    add_center(doc, DATE_TEXT, 14, False, "宋体")
    add_page_break(doc)


def write_reference_verification():
    references = [
        {
            "title": "Qt 6 Documentation: Qt Widgets and Qt Network",
            "authors": "The Qt Company",
            "year": 2026,
            "source": "Qt Documentation",
            "doi_or_url": "https://doc.qt.io/",
            "relevance_note": "用于说明Qt Widgets界面开发和Qt网络通信模块。",
            "status": "verified_by_official_documentation",
        },
        {
            "title": "MySQL 8.4 Reference Manual",
            "authors": "Oracle Corporation",
            "year": 2026,
            "source": "MySQL Documentation",
            "doi_or_url": "https://dev.mysql.com/doc/",
            "relevance_note": "用于说明MySQL关系数据库、表结构和事务处理。",
            "status": "verified_by_official_documentation",
        },
        {
            "title": "CMake Documentation",
            "authors": "Kitware",
            "year": 2026,
            "source": "CMake Documentation",
            "doi_or_url": "https://cmake.org/documentation/",
            "relevance_note": "用于说明项目构建工具和跨平台工程组织。",
            "status": "verified_by_official_documentation",
        },
        {
            "title": "Software Engineering at Google",
            "authors": "Titus Winters, Tom Manshreck, Hyrum Wright",
            "year": 2020,
            "source": "O'Reilly Media",
            "doi_or_url": "ISBN 978-1492082798",
            "relevance_note": "用于支撑工程组织、代码维护和测试验证等软件工程内容。",
            "status": "bibliographic_reference",
        },
        {
            "title": "Microsoft C++ Documentation",
            "authors": "Microsoft",
            "year": 2026,
            "source": "Microsoft Learn",
            "doi_or_url": "https://learn.microsoft.com/cpp/",
            "relevance_note": "用于说明C++工程开发、类设计和跨平台编译相关内容。",
            "status": "verified_by_official_documentation",
        },
        {
            "title": "NIST SP 800-53 Rev. 5 Security and Privacy Controls",
            "authors": "National Institute of Standards and Technology",
            "year": 2020,
            "source": "NIST",
            "doi_or_url": "https://doi.org/10.6028/NIST.SP.800-53r5",
            "relevance_note": "用于说明访问控制、审计和安全控制思想。",
            "status": "standards_reference",
        },
    ]
    ACTIVE_REFERENCES_OUTPUT.write_text(json.dumps(references, ensure_ascii=False, indent=2), encoding="utf-8")


def write_summary_markdown():
    ACTIVE_SUMMARY_MD.write_text(
        "# 基于Qt和MySQL的医院门诊挂号与药品管理系统课程设计总结\n\n"
        "本文件为成稿生成过程的源稿摘要，完整格式化正文以DOCX为准。\n\n"
        "## 主要交付物\n\n"
        "- 医院门诊挂号与药品管理系统-课程设计总结.docx\n"
        "- 医院门诊挂号与药品管理系统-课程设计总结-附件.docx\n"
        "- 医院门诊挂号与药品管理系统-课程设计总结-文献核验清单.json\n\n"
        "## 图表\n\n"
        "图3-1至图3-8覆盖系统架构、功能架构、技术架构、类图、时序图和E-R图；"
        "图4-1至图4-4覆盖目录结构、权限流程、处方发药流程和收费退费流程。\n",
        encoding="utf-8",
    )


def build_appendix(images: dict[str, Path]):
    doc = Document(ACTIVE_TEMPLATE)
    clear_body(doc)
    configure_document(doc)
    add_heading(doc, "图表与关键设计附件", 1)
    add_paragraph(doc, "本附件用于收录课程设计总结中主要图表的设计说明和关键源码摘录，便于后续修改图形或答辩说明。")
    add_table(doc, "表A-1 图表清单", ["图号", "图名", "文件"], [
        ["图3-1", "系统总体架构图", str(images["arch"].relative_to(ROOT))],
        ["图3-2", "系统功能架构图", str(images["function"].relative_to(ROOT))],
        ["图3-3", "技术架构图", str(images["tech"].relative_to(ROOT))],
        ["图3-4", "初步类图设计", str(images["class_initial"].relative_to(ROOT))],
        ["图3-5", "挂号业务时序图", str(images["sequence"].relative_to(ROOT))],
        ["图3-6", "详细类图设计", str(images["class_detail"].relative_to(ROOT))],
        ["图3-7", "数据库实体关系图", str(images["er"].relative_to(ROOT))],
        ["图3-8", "核心业务E-R图", str(images["core_er"].relative_to(ROOT))],
        ["图4-1", "程序目录结构图", str(images["directory"].relative_to(ROOT))],
        ["图4-2", "登录权限控制流程图", str(images["permission_flow"].relative_to(ROOT))],
        ["图4-3", "处方审核与发药流程图", str(images["prescription_flow"].relative_to(ROOT))],
        ["图4-4", "收费结算与退费流程图", str(images["payment_flow"].relative_to(ROOT))],
        ["图4-5", "平台端主界面布局图", str(images["platform_ui"].relative_to(ROOT))],
        ["图4-6", "患者端预约界面布局图", str(images["patient_ui"].relative_to(ROOT))],
    ])
    add_heading(doc, "A.1 核心E-R图关系说明", 2)
    add_paragraph(doc, "患者、医生、排班、挂号、病历、处方、药品、账单和支付构成系统核心业务链路。挂号记录同时关联患者、医生和排班，接诊后继续关联病历、检查、处方和账单。处方明细与药品表关联，发药或退药时写入库存流水。")
    add_heading(doc, "A.2 核心接口摘录", 2)
    add_code_listing(doc, "附件文件A-1 请求路由接口", """
void RequestRouter::registerService(const QString& module, ModuleService* service);
common::Response RequestRouter::route(const common::Request& request) const;
bool RequestRouter::isAllowed(const QString& roleCode,
                              const QString& module,
                              const QString& action) const;
""")
    add_code_listing(doc, "附件文件A-2 处方流程动作", """
createPrescription -> reviewPrescription -> payBill -> dispensePrescription
returnPrescription -> stock_records回补 -> operation_logs记录
""")
    doc.save(ACTIVE_APPENDIX_OUTPUT)


def build_document():
    images = generate_images()
    write_reference_verification()
    write_summary_markdown()
    build_appendix(images)
    doc = Document(ACTIVE_TEMPLATE)
    clear_body(doc)
    configure_document(doc)
    build_cover(doc)

    add_heading(doc, "摘  要", 1)
    add_paragraph(
        doc,
        "随着医院门诊业务量增长，传统人工挂号、纸质病历和分散收费方式在排队效率、信息共享、库存核对和费用统计等方面存在明显不足。"
        "本课程设计围绕医院门诊挂号与药品管理场景，设计并实现了一套基于Qt和MySQL的医院门诊挂号与药品管理系统。"
        "系统采用客户端/服务端分离架构，客户端使用Qt Widgets构建桌面端工作台和患者预约入口，服务端基于Qt TCP通信提供统一业务接口，"
        "客户端与服务端之间通过JSON Lines格式传输请求和响应，数据库采用MySQL保存患者、科室、医生、排班、挂号、病历、检查、处方、药品、库存、账单、支付、权限和日志等数据。"
        "系统实现了登录鉴权、患者管理、预约挂号、候诊叫号、医生接诊、检查检验、处方审核发药、药品库存预警、收费结算、费用统计、权限配置和操作日志等功能。"
        "通过模块化页面、统一路由、会话校验和事务化数据处理，系统能够较完整地支撑门诊核心业务流程，满足课程设计对功能完整性、数据库设计和工程实现的要求。",
    )
    add_paragraph(doc, "关键词：医院门诊；挂号管理；药品管理；Qt；MySQL；C/S架构", first_line=False)
    add_page_break(doc)
    add_toc(doc)

    add_heading(doc, "第一章  绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_paragraph(doc, "门诊是医院对外服务中业务频率最高、数据流转最密集的环节之一。患者从建档、挂号、候诊、接诊到检查、处方、收费和取药，需要多个岗位连续协作。如果仍然依赖人工登记和分散表格，容易出现患者排队时间长、医生号源更新不及时、药品库存不准确、收费状态难追踪等问题。")
    add_paragraph(doc, "本项目以医院门诊常见业务为对象，选择Qt作为客户端和服务端开发框架，选择MySQL作为关系数据库，围绕挂号与药品管理主线完成系统设计与实现。Qt具备跨平台桌面界面能力和网络通信能力，MySQL适合保存结构化业务数据，两者结合能够满足课程设计对可运行系统、数据库应用和模块化工程的要求。")
    add_paragraph(doc, "从实际业务看，门诊系统的难点并不只在单个页面的增删改查，而在于业务状态跨岗位流转。挂号员需要看到医生排班和剩余号源，医生需要看到候诊患者和既往病历，药师需要根据处方状态和缴费状态决定是否发药，收费员需要处理挂号费、药品费、检查费和退费请求。任何一个环节的数据不同步，都可能影响患者就诊体验。因此，本项目在设计时没有把各页面做成孤立模块，而是围绕患者、号源、挂号单、病历、处方、账单和库存形成连续的数据链路。")
    add_heading(doc, "1.1.1 国内研究现状", 3)
    add_paragraph(doc, "国内医院信息化建设已逐步从单一收费、挂号系统发展到覆盖预约、诊疗、药房、医保和运营统计的一体化平台。多数大型医院已部署HIS、EMR、PACS等系统，但基层或教学场景中的轻量级门诊系统仍需要在功能完整、部署简单、数据结构清晰之间取得平衡。对于课程设计而言，围绕门诊业务构建一个结构清楚、流程闭环的系统，有助于理解医院信息系统的基本组织方式。")
    add_paragraph(doc, "在教学和中小型管理系统场景中，常见实现方式包括桌面客户端、Web后台和移动端预约入口。桌面客户端适合固定岗位使用，输入效率较高；Web系统部署方便，适合多终端访问；移动端更适合患者预约和结果查询。本项目结合课程设计要求和现有技术条件，采用Qt桌面客户端作为主要交互入口，同时在客户端中加入患者预约窗口，使医院端和患者端能够共用同一套服务端接口。")
    add_heading(doc, "1.1.2 国外研究现状", 3)
    add_paragraph(doc, "国外医疗信息系统更强调预约分诊、电子病历互通、患者自助服务和数据安全。相关系统通常采用分层架构、标准化接口和权限审计机制，以降低不同系统之间的数据交换成本。本项目虽然规模较小，但在设计中引入了客户端/服务端分离、统一通信协议、权限校验、业务日志和状态流转等思想，使系统具备一定的工程扩展基础。")
    add_paragraph(doc, "国外门诊系统在工程实践中普遍重视数据标准、访问控制和审计追踪。虽然本课程设计没有接入真实医疗标准接口，但在实现中保留了权限表、操作日志表、审计明细表和医保模拟交易表等结构，能够体现医疗业务系统对身份、流程和数据留痕的基本要求。")
    add_heading(doc, "1.2 研究的目的与意义", 2)
    add_heading(doc, "1.2.1 研究的目的", 3)
    add_paragraph(doc, "本课程设计的主要目的是完成一个可运行、可演示、结构清晰的医院门诊挂号与药品管理系统。系统需要覆盖患者建档、医生排班、预约挂号、候诊接诊、检查处方、药品库存、收费结算和统计审计等流程，并通过数据库保存关键业务数据。")
    add_paragraph(doc, "在技术目标上，系统需要验证Qt客户端界面开发、TCP网络通信、JSON协议封装、MySQL数据库操作、角色权限控制、业务状态流转和跨平台构建脚本等能力。报告撰写时也需要将这些实现细节转化为需求分析、系统设计、数据库设计和系统实现等文档内容，形成完整的软件工程交付材料。")
    add_heading(doc, "1.2.2 研究的意义", 3)
    add_paragraph(doc, "从学习意义看，项目综合运用了C++、Qt界面开发、TCP网络通信、JSON数据交换、MySQL数据库设计、CMake工程组织和测试验证等知识。从应用意义看，系统将门诊业务拆分为多个可维护模块，能够提高挂号与药品管理流程的规范性，也为后续扩展医保结算、移动预约和数据分析提供基础。")
    add_paragraph(doc, "从工程训练角度看，本项目不是单纯完成界面展示，而是通过客户端、服务端、数据库和测试文件共同构成一个可运行项目。开发过程中需要考虑不同角色看到的页面不同、接口不能绕过权限直接访问、挂号不能重复占用号源、处方发药需要扣减库存、支付和退费需要改变账单状态等问题。这些问题能够帮助学生从功能实现进一步过渡到业务规则实现。")
    add_heading(doc, "1.3 系统要解决的主要问题及论文结构", 2)
    add_heading(doc, "1.3.1 系统要完成的主要功能及描述", 3)
    add_paragraph(doc, "系统围绕门诊业务闭环设计功能：登录权限模块负责用户认证、角色权限和会话管理；患者管理模块负责患者档案维护；挂号排班模块负责科室、医生、号源和预约挂号；诊疗模块负责候诊队列、医生接诊、病历和检查；处方药品模块负责处方审核、发药、退药和库存预警；收费统计模块负责账单、支付、退费和收入统计；系统管理模块负责权限配置和操作日志。")
    add_heading(doc, "1.3.2 论文结构", 3)
    add_paragraph(doc, "本文第一章介绍研究背景、目的意义和主要功能；第二章分析系统的总体需求、功能需求和非功能需求；第三章完成系统总体设计、类图设计、时序图设计和数据库设计；第四章说明开发环境、程序目录结构和核心模块实现；第五章总结课程设计成果并提出后续改进方向。")

    add_heading(doc, "第二章 需求分析", 1)
    add_heading(doc, "2.1 系统总体需求分析", 2)
    add_paragraph(doc, "系统面向医院门诊业务人员和患者用户。医院端用户包括管理员、院长、挂号员、医生、药师和收费员，不同角色进入系统后只能看到对应功能模块。患者端用户可以维护个人信息并完成预约挂号。系统需要通过服务端统一处理业务请求，保证数据访问、权限判断和操作日志集中管理。")
    add_paragraph(doc, "总体需求可以概括为“一套数据、多端入口、分角色处理、全过程留痕”。一套数据指患者、科室、医生、排班、挂号、病历、处方、库存和账单统一存储在MySQL数据库中；多端入口指医院工作人员端和患者预约入口通过统一服务端进行交互；分角色处理指系统按管理员、院长、挂号员、医生、药师、收费员和患者划分功能边界；全过程留痕指关键业务操作写入操作日志和审计明细，便于后期追溯。")
    add_table(doc, "表2-1 用户角色与主要需求", ["角色", "主要需求"], [
        ["管理员", "维护用户、角色、权限、科室、医生和系统基础数据，查看操作日志。"],
        ["挂号员", "维护患者档案，查询号源，完成挂号、退号和候诊队列管理。"],
        ["医生", "查看候诊患者，完成叫号、接诊、病历记录、检查申请和处方开立。"],
        ["药师", "审核处方、发药、退药，维护药品库存和库存预警。"],
        ["收费员", "处理账单支付、退费审核和费用统计。"],
        ["患者", "注册登录、维护个人信息、选择科室医生并预约挂号。"],
    ])
    add_heading(doc, "2.2 系统功能需求分析", 2)
    add_paragraph(doc, "系统功能需求以门诊业务流程为主线，既包含基础数据维护，也包含跨模块的业务状态流转。挂号成功后生成挂号记录和账单，候诊队列可根据状态展示等待、已叫号、接诊中和已完成患者；医生接诊后可以保存病历、开立检查和处方；处方经过审核、缴费后进入发药流程；收费和库存变更均需要形成可追溯记录。")
    add_paragraph(doc, "功能设计时需要保证模块之间既相对独立又能通过业务编号关联。例如患者管理模块主要维护患者基础资料，但挂号模块需要引用患者编号；医生排班模块维护号源，挂号模块需要扣减剩余号源；医生接诊模块生成病历和处方，药房模块需要读取处方明细并改变处方状态；收费模块处理账单，统计模块需要根据支付记录形成日报数据。")
    add_table(doc, "表2-2 核心功能需求", ["模块", "功能说明"], [
        ["登录权限", "根据用户名、密码和角色建立会话Token，并控制模块可见性和接口访问权限。"],
        ["患者管理", "支持患者新增、查询、修改和患者端账号关联。"],
        ["挂号管理", "支持科室医生联动、号源选择、医保校验、挂号、退号和费用生成。"],
        ["医生排班", "维护医生出诊日期、时段、总号源和剩余号源，支持智能排班规则。"],
        ["医生接诊", "管理候诊队列、叫号、开始接诊、填写诊断、申请检查和开立处方。"],
        ["药品库存", "维护药品分类、条码、价格、库存数量、有效期和预警数量。"],
        ["收费结算", "支持自费/医保支付、支付状态查询、退费申请和退费审核。"],
        ["统计审计", "展示收入、挂号量、医生工作量、库存预警和操作日志。"],
    ])
    add_table(doc, "表2-3 关键业务状态流转", ["业务对象", "主要状态", "触发模块"], [
        ["挂号记录", "待支付、待就诊、已叫号、接诊中、已完成、已退号", "挂号管理、候诊队列、医生接诊、收费结算"],
        ["医生排班", "正常、停诊、剩余号源变化", "医生排班、挂号管理"],
        ["处方", "待审核、已审核、已驳回、已发药、已退药", "医生接诊、处方管理、药品库存"],
        ["账单", "待支付、已支付、退费申请中、已退费、已取消", "收费结算、挂号管理、处方管理"],
        ["库存记录", "入库、出库、发药扣减、退药回补、库存预警", "药品库存、处方管理"],
    ])
    add_heading(doc, "2.3 系统非功能需求分析", 2)
    add_paragraph(doc, "系统非功能需求主要包括易用性、可靠性、安全性、可维护性和可移植性。易用性方面，客户端采用左侧模块导航和统一表格操作，便于不同岗位快速定位功能；可靠性方面，服务端集中处理数据库连接和异常响应；安全性方面，通过会话Token、角色权限和操作日志降低越权和误操作风险；可维护性方面，客户端页面、服务端业务服务和数据库脚本按模块组织；可移植性方面，工程兼容Windows、Linux和银河麒麟环境。")
    add_paragraph(doc, "可靠性方面，服务端数据库连接由DatabaseManager集中管理，查询和写入前会检查连接状态，必要时尝试重新连接；演示模式下也保留DemoRepository作为数据回退方案。安全性方面，客户端虽然会根据角色隐藏不可用模块，但真正的权限控制仍由服务端RequestRouter执行，避免只依赖界面隐藏。可维护性方面，服务端业务模块统一继承ModuleService并实现handle方法，新增模块时只需要注册到路由器并补充相应页面。")
    add_table(doc, "表2-4 非功能需求与实现约束", ["需求类型", "实现约束"], [
        ["易用性", "页面保持统一表格、筛选、刷新和操作按钮布局，减少不同岗位学习成本。"],
        ["可靠性", "数据库连接集中管理，核心写操作使用事务或状态校验，失败时返回明确错误信息。"],
        ["安全性", "登录后发放Token，服务端根据Token、角色和模块动作进行访问控制。"],
        ["可维护性", "客户端页面、服务端服务、数据库表和测试文件按业务模块组织。"],
        ["可移植性", "使用CMake组织工程，提供Windows、Linux和银河麒麟构建运行脚本。"],
        ["可测试性", "通过协议测试、权限路由测试、排班规则测试和业务源码测试验证关键逻辑。"],
    ])
    add_heading(doc, "2.4 本章小结", 2)
    add_paragraph(doc, "本章从用户角色、业务流程、功能模块和非功能约束四个方面分析了系统需求。分析结果表明，系统需要以挂号、接诊、处方、库存和收费为核心，配套权限、统计和审计能力，并采用清晰的分层架构保证后续实现。")

    add_heading(doc, "第三章 系统设计", 1)
    add_heading(doc, "3.1 系统总体设计", 2)
    add_heading(doc, "3.1.1 系统架构", 3)
    add_paragraph(doc, "系统采用客户端/服务端架构。客户端负责界面展示、表单校验和用户交互，服务端负责业务路由、权限校验、数据库访问和日志记录。客户端与服务端通过TCP长连接通信，请求与响应均采用JSON Lines格式，便于按行解析和调试。")
    add_paragraph(doc, "系统架构设计遵循职责分离原则。客户端不直接连接数据库，而是通过ApiClient封装请求并提交给服务端；服务端接收请求后先完成协议解析和会话校验，再根据module和action定位具体业务服务；业务服务只处理本模块规则，并通过DatabaseManager访问数据库。这样做可以避免客户端分散保存SQL语句，也便于统一处理权限、异常和日志。")
    add_figure(doc, images["arch"], "图3-1 系统总体架构图")
    add_heading(doc, "3.1.2 功能架构", 3)
    add_paragraph(doc, "系统功能按业务职责划分为基础管理、挂号排班、诊疗业务、药品收费和统计审计五类。各模块之间通过服务端统一接口进行数据交互，避免客户端直接操作数据库。")
    add_paragraph(doc, "功能架构中，基础管理提供用户、角色、科室、医生和患者档案等基础数据；挂号排班负责从医生号源到挂号记录的生成；诊疗业务负责候诊、接诊、病历、检查和处方；药品收费负责处方审核、发药扣库存、账单支付和退费；统计审计负责将业务结果转换为管理视图和追踪记录。")
    add_figure(doc, images["function"], "图3-2 系统功能架构图", 15.2)
    add_heading(doc, "3.1.3 技术架构", 3)
    add_paragraph(doc, "技术架构分为表现层、通信层、服务层、业务层、数据层和工程支撑层。表现层使用Qt Widgets，通信层使用QTcpSocket和JSON协议，服务层由HospitalServer、ClientConnection和RequestRouter组成，业务层由多个ModuleService实现，数据层以MySQL为主并预留Redis能力。")
    add_table(doc, "表3-1 通信请求格式设计", ["字段", "含义", "示例"], [
        ["module", "业务模块名称", "registration、billing、prescription"],
        ["action", "模块内动作名称", "list、create、pay、review"],
        ["headers", "会话Token、角色等头部信息", "token、roleCode"],
        ["payload", "业务参数对象", "patientId、doctorId、scheduleId"],
    ])
    add_figure(doc, images["tech"], "图3-3 技术架构图")
    add_heading(doc, "3.2 系统类图设计", 2)
    add_heading(doc, "3.2.1 初步类图设计", 3)
    add_paragraph(doc, "初步类图体现系统的主要对象和调用关系。MainWindow负责模块导航，ModulePage提供通用页面能力，ApiClient封装网络请求，服务端通过HospitalServer接收连接，再由RequestRouter转发到各业务服务，DatabaseManager负责数据库访问。")
    add_paragraph(doc, "客户端页面大多继承ModulePage，该基类封装了表格、刷新、请求发送和响应处理等通用逻辑。不同业务页面只需要补充本模块字段、按钮和特殊操作。服务端则通过ModuleService定义统一服务接口，保证RequestRouter能够以一致方式调用各个业务模块。")
    add_figure(doc, images["class_initial"], "图3-4 初步类图设计", 15.2)
    add_heading(doc, "3.2.2 时序图设计", 3)
    add_paragraph(doc, "挂号流程是系统的核心流程之一。用户在挂号页面选择患者、科室、医生和号源后，客户端封装请求发送至服务端，服务端完成权限校验、号源校验、医保状态校验、事务写入和账单生成，最后返回结果并刷新客户端列表。")
    add_paragraph(doc, "时序设计中特别强调号源校验和事务处理。挂号请求不能只新增一条挂号记录，还需要检查患者身份、医生排班、剩余号源、医保预校验状态，并在成功后同步更新排班剩余号源和账单信息。如果其中任一环节失败，服务端应返回失败响应，客户端保持原列表并提示用户重新选择。")
    add_figure(doc, images["sequence"], "图3-5 挂号业务时序图", 15.2)
    add_heading(doc, "3.2.3 详细类图设计", 3)
    add_paragraph(doc, "详细类图突出服务端业务服务的模块化设计。RequestRouter负责统一入口，各业务服务继承ModuleService并实现handle方法，业务服务共享DatabaseManager进行数据库访问。这样的设计便于新增模块，也便于对权限、日志和响应格式进行统一处理。")
    add_table(doc, "表3-2 服务端主要类职责", ["类名", "职责"], [
        ["HospitalServer", "监听端口，接收客户端TCP连接。"],
        ["ClientConnection", "按行读取请求，调用路由器并返回JSON响应。"],
        ["RequestRouter", "解析模块动作，校验会话权限，转发业务服务，记录操作日志。"],
        ["SessionManager", "创建、查询和失效登录会话Token。"],
        ["DatabaseManager", "管理数据库连接、兼容性建表、事务与查询。"],
        ["ModuleService", "业务服务基类，定义统一handle接口。"],
    ])
    add_figure(doc, images["class_detail"], "图3-6 详细类图设计", 15.2)
    add_heading(doc, "3.3 系统数据库设计", 2)
    add_heading(doc, "3.3.1 数据库实体模型设计", 3)
    add_paragraph(doc, "数据库以门诊业务实体为核心建立关系模型。患者、医生、科室和排班构成挂号基础；挂号记录连接病历、检查、处方和账单；处方明细与药品表关联，库存流水记录药品数量变化；支付表与账单关联，操作日志记录关键业务动作。")
    add_figure(doc, images["er"], "图3-7 数据库实体关系图", 15.5)
    add_paragraph(doc, "由于数据库总表较多，完整E-R图用于说明整体数据域，核心业务E-R图则重点展示患者从挂号到处方、账单和支付的主链路。权限、医保、审计和出库等扩展表不直接放入核心图，以避免核心业务关系被过多辅助表淹没。")
    add_figure(doc, images["core_er"], "图3-8 核心业务E-R图", 15.2)
    add_heading(doc, "3.3.2 数据库设计", 3)
    add_paragraph(doc, "系统数据库名为hospital_outpatient，建库脚本位于database/schema.sql。脚本可重复执行，会重建业务表并写入演示数据。核心表通过主键、唯一键、外键和索引约束保证数据一致性。")
    add_table(doc, "表3-3 主要数据库表设计", ["表名", "用途", "关键字段"], [
        ["users / roles", "医院工作人员账号与角色", "username、password_hash、role_id、role_code"],
        ["patients / patient_users", "患者档案和患者端账号", "patient_no、name、id_card、phone、user_id"],
        ["departments / doctors", "科室与医生基础数据", "dept_code、dept_name、doctor_id、registration_fee"],
        ["doctor_schedules", "医生排班和号源", "doctor_id、work_date、period、total_quota、remain_quota"],
        ["registrations", "挂号记录", "registration_no、patient_id、doctor_id、schedule_id、status、fee"],
        ["medical_records / examinations", "病历与检查检验", "registration_id、diagnosis、item_name、status"],
        ["prescriptions / prescription_items", "处方主表与处方明细", "prescription_no、drug_id、quantity、amount、status"],
        ["drugs / stock_records", "药品资料与库存流水", "drug_code、barcode、stock_quantity、warning_quantity、change_type"],
        ["bills / payments", "账单与支付记录", "bill_no、total_amount、status、payment_no、pay_method"],
        ["operation_logs / audit_log_details", "操作日志和审计明细", "module、action、business_key、old_value、new_value"],
    ])
    add_table(doc, "表3-4 挂号主表字段设计", ["字段名", "类型含义", "说明"], [
        ["registration_id", "主键", "挂号记录唯一标识。"],
        ["registration_no", "业务编号", "挂号单号，用于页面展示和业务追踪。"],
        ["patient_id", "外键", "关联患者档案。"],
        ["doctor_id", "外键", "关联接诊医生。"],
        ["schedule_id", "外键", "关联医生排班号源。"],
        ["status", "状态字段", "表示待支付、待就诊、已叫号、接诊中、已完成等状态。"],
        ["fee", "金额字段", "保存挂号费用。"],
    ])
    add_heading(doc, "3.4 本章小结", 2)
    add_paragraph(doc, "本章完成了系统的总体设计、功能设计、技术设计、类设计、时序设计和数据库设计。设计结果表明，系统采用分层和模块化方式组织，核心业务关系清晰，能够支撑第四章的具体实现。")

    add_heading(doc, "第四章 系统实现", 1)
    add_heading(doc, "4.1 系统开发环境", 2)
    add_table(doc, "表4-1 系统开发环境", ["类别", "内容"], [
        ["开发语言", "C++"],
        ["界面框架", "Qt Widgets，兼容Qt5/Qt6"],
        ["服务端框架", "Qt TCP Server、QTcpSocket、JSON处理组件"],
        ["数据库", "MySQL，Windows可通过QODBC连接，Linux可通过QMYSQL连接"],
        ["构建工具", "CMake、CMakePresets、Qt Creator、VS Code"],
        ["运行环境", "Windows、Linux、银河麒麟桌面环境"],
        ["测试方式", "协议测试、权限路由测试、业务源码测试、数据库兼容性测试"],
    ])
    add_paragraph(doc, "本项目采用CMake组织多目标工程，common、server、client和launcher分别作为独立构建单元。Windows环境下使用Qt 6 MinGW套件进行开发和运行，Linux与银河麒麟环境下通过脚本安装依赖、配置构建目录并启动服务端和客户端。数据库连接在Windows下可以通过QODBC连接MySQL，在Linux下可以通过QMYSQL驱动连接MySQL。")
    add_heading(doc, "4.2 程序目录结构", 2)
    add_paragraph(doc, "项目采用清晰的工程目录组织方式。client目录保存客户端界面和页面模块，server目录保存服务端路由、业务服务和数据库访问，common目录保存客户端和服务端共享协议，database目录保存建库脚本，scripts目录保存跨平台构建与运行脚本，tests目录保存自动化测试。")
    add_figure(doc, images["directory"], "图4-1 程序目录结构图")
    add_heading(doc, "4.3 平台端", 2)
    add_paragraph(doc, "平台端即医院工作人员使用的桌面客户端。MainWindow负责主窗口、左侧导航和模块权限过滤，按照用户角色加载院长驾驶舱、患者管理、挂号管理、候诊队列、患者病历档案、科室管理、医生排班、医生管理、医生接诊、检查检验、处方管理、药品库存、收费结算、费用统计、操作日志和权限配置等页面。")
    add_figure(doc, images["platform_ui"], "图4-5 平台端主界面布局图", 15.2)
    add_paragraph(doc, "服务端实现中，HospitalServer负责监听客户端连接，ClientConnection负责按行接收JSON请求，RequestRouter负责解析module和action并完成会话鉴权、权限判断、业务转发和操作日志写入。各业务模块以ModuleService子类形式实现，能够保持接口统一并降低模块之间耦合。")
    add_paragraph(doc, "挂号管理模块支持科室、医生、号源联动和挂号保存；医生排班模块支持按日期、科室、医生维护号源；医生接诊模块支持候诊队列、叫号、接诊、病历、检查和处方；药品库存模块支持药品信息、库存流水和预警；收费结算模块支持自费、医保模拟支付、支付状态、退费申请和审核。")
    add_heading(doc, "4.3.1 登录权限模块实现", 3)
    add_paragraph(doc, "登录权限模块由客户端登录窗口、服务端AuthService、SessionManager、AuthorizationService和RequestRouter共同完成。用户登录成功后，服务端返回Token、角色编码和权限集合；客户端根据权限渲染模块入口；后续请求仍由服务端根据Token和module/action进行二次校验。这样的设计避免了只依赖客户端隐藏按钮造成的越权风险。")
    add_figure(doc, images["permission_flow"], "图4-2 登录权限控制流程图", 15.2)
    add_code_listing(doc, "文件4-1 服务端统一路由接口片段", """
class ModuleService
{
public:
    virtual ~ModuleService() = default;
    virtual common::Response handle(const common::Request& request) = 0;
};

class RequestRouter
{
public:
    void registerService(const QString& module, ModuleService* service);
    common::Response route(const common::Request& request) const;
private:
    bool isAllowed(const QString& roleCode,
                   const QString& module,
                   const QString& action) const;
};
""")
    add_paragraph(doc, "从代码结构可以看出，业务服务不是由客户端直接调用，而是统一注册到RequestRouter。路由器在route方法中负责会话查找、权限判断、业务分发和响应包装。权限配置模块还可以维护角色与菜单之间的关系，使系统能够根据不同岗位调整功能访问范围。")

    add_heading(doc, "4.3.2 挂号与排班模块实现", 3)
    add_paragraph(doc, "挂号与排班模块是系统的业务起点。排班页面负责维护医生出诊日期、时段、总号源和剩余号源，并提供智能排班、未排班医生提示和号源重置能力。挂号页面根据患者、科室、医生和排班信息创建挂号记录，支持医保预校验、急诊标记、退号和候诊状态刷新。服务端RegistrationService在创建挂号时需要同时处理患者身份、排班剩余号源、重复挂号、医保状态和账单生成。")
    add_table(doc, "表4-2 挂号与排班模块核心实现", ["实现点", "说明"], [
        ["科室医生联动", "客户端根据科室筛选医生，再根据医生与日期筛选排班号源。"],
        ["号源扣减", "挂号成功后更新doctor_schedules.remain_quota。"],
        ["医保预校验", "根据参保类型、异地备案和欠费状态决定是否允许医保挂号。"],
        ["候诊状态", "挂号记录进入待就诊队列，医生或挂号员可以叫号和标记急诊。"],
        ["账单生成", "挂号成功后生成对应账单，收费模块继续处理支付。"],
    ])
    add_code_listing(doc, "文件4-2 排班页面关键成员片段", """
class SchedulePage : public ModulePage
{
private:
    QJsonArray loadServerScheduleRange(const QDate& startDate,
                                       const QDate& endDate,
                                       bool* ok = nullptr) const;
    bool hasActiveClinicCoverage(const QJsonArray& rows,
                                 const QString& clinic,
                                 const QString& date) const;
    bool shouldDoctorWorkOnRotation(const QString& doctor,
                                    int dayOffset,
                                    const QStringList& clinicDoctors) const;
};
""")
    add_paragraph(doc, "排班页面中的loadServerScheduleRange用于加载指定日期范围内的排班数据，hasActiveClinicCoverage用于判断某科室某日是否已有有效出诊安排，shouldDoctorWorkOnRotation用于根据轮转规则判断医生是否应排班。这些函数体现了排班模块不只是保存表格数据，还包含一定的业务规则。")

    add_heading(doc, "4.3.3 医生接诊与病历模块实现", 3)
    add_paragraph(doc, "医生接诊模块以候诊队列为入口。医生可以查看等待患者、叫号、开始接诊，并在接诊过程中填写主诉、现病史、既往史、体格检查、诊断和医嘱等内容。系统还支持检查项目申请和处方开立，病历修改会写入审计明细，便于追踪重要字段变化。")
    add_table(doc, "表4-3 医生接诊模块功能点", ["功能点", "实现说明"], [
        ["候诊队列", "WaitingQueuePage根据挂号状态统计等待、已叫号、急诊和平均等待时间。"],
        ["叫号接诊", "RegistrationService提供call、start等动作，改变挂号状态。"],
        ["病历维护", "PatientRecordService保存病历字段并生成审计明细。"],
        ["检查申请", "ExaminationService维护检查项目和检查状态。"],
        ["处方开立", "PrescriptionService创建处方主表和明细表。"],
    ])
    add_code_listing(doc, "文件4-3 医生接诊页面关键成员片段", """
class ConsultationPage : public ModulePage
{
private slots:
    void callSelectedPatient();
    void startConsultation();
    void onConsultationResponse(const common::Response& response);
private:
    void loadExaminationItems();
    QJsonObject m_pendingExamRequest;
    QJsonObject m_pendingPrescriptionRequest;
};
""")
    add_paragraph(doc, "接诊页面中保留m_pendingExamRequest和m_pendingPrescriptionRequest两个临时对象，用于在医生填写接诊信息时暂存检查和处方相关请求。这样可以把一次接诊中的病历、检查和处方动作组织在同一业务场景下，减少页面跳转和重复录入。")

    add_heading(doc, "4.3.4 处方药品模块实现", 3)
    add_paragraph(doc, "处方药品模块连接医生、药师、收费和库存。医生开立处方后，处方进入待审核状态；药师审核时可以依据pass_rules进行规则提示；收费完成后，药师执行发药操作，系统扣减药品库存并写入库存流水；如果发生退药，系统需要回补库存并改变处方状态。")
    add_figure(doc, images["prescription_flow"], "图4-3 处方审核与发药流程图", 15.2)
    add_code_listing(doc, "文件4-4 处方服务动作片段", """
common::Response PrescriptionService::handle(const common::Request& request)
{
    const QString action = request.action;
    if (action == "create") return createPrescriptionInDatabase(m_database, request.payload);
    if (action == "review") return reviewPrescriptionInDatabase(m_database, request.payload);
    if (action == "dispense") return dispensePrescriptionInDatabase(m_database, request.payload);
    if (action == "return") return returnPrescriptionInDatabase(m_database, request.payload);
    return {false, "未知处方操作。", {}};
}
""")
    add_paragraph(doc, "处方服务按动作拆分为创建、审核、发药和退药。创建处方时写入prescriptions和prescription_items；审核时改变处方状态并记录审核人；发药时校验支付状态和库存数量；退药时写入库存回补记录。通过服务端集中处理这些动作，可以避免不同页面对处方状态理解不一致。")

    add_heading(doc, "4.3.5 收费结算与统计审计模块实现", 3)
    add_paragraph(doc, "收费结算模块处理挂号费、检查费和药品费等账单。系统支持自费支付、医保模拟支付、扫码支付状态查询、退费申请和退费审核。统计模块根据费用与挂号数据展示日收入、科室收入、医生工作量和库存预警，院长驾驶舱则用于汇总关键指标。")
    add_figure(doc, images["payment_flow"], "图4-4 收费结算与退费流程图", 15.2)
    add_code_listing(doc, "文件4-5 收费服务接口片段", """
class BillingService : public ModuleService
{
private:
    common::Response processSelfPay(const QJsonObject& payload);
    common::Response processMedicalInsurancePay(const QJsonObject& payload);
    common::Response createPaymentQr(const QJsonObject& payload);
    common::Response checkPayStatus(const QJsonObject& payload);
    common::Response requestRefundBill(const QJsonObject& payload);
    common::Response reviewRefundBill(const QJsonObject& payload);
};
""")
    add_paragraph(doc, "收费模块的实现重点在于账单状态和支付记录的一致性。支付成功后需要同时更新bills和payments，退费申请后不能立即改变为已退费，而是进入审核流程；审核通过后再写入退费结果和审计日志。这样的状态拆分能够更接近真实收费业务。")

    add_heading(doc, "4.3.6 权限配置与操作日志模块实现", 3)
    add_paragraph(doc, "权限配置模块用于维护用户、角色和权限项。管理员可以创建用户、重置密码、启停账号以及为角色勾选功能权限。操作日志模块记录模块名、动作、业务键、操作人和操作结果，审计明细表保存关键字段变更前后的值。")
    add_table(doc, "表4-4 测试与验证文件", ["测试文件", "验证内容"], [
        ["protocol_tests.cpp", "验证请求响应JSON协议编码和解码。"],
        ["auth_router_tests.cpp", "验证路由器鉴权和模块访问控制。"],
        ["schedule_rule_engine_tests.cpp", "验证排班规则引擎。"],
        ["payment_router_tests.cpp", "验证收费路由与支付状态处理。"],
        ["registration_insurance_source_tests.cpp", "验证挂号医保预校验相关源码规则。"],
        ["pharmacy_workflow_source_tests.cpp", "验证处方审核、发药和库存流程源码。"],
    ])
    add_heading(doc, "4.4 患者端", 2)
    add_paragraph(doc, "项目中患者端以PatientLoginDialog和PatientAppointmentWindow实现，承担患者注册登录、个人信息维护和预约挂号功能。患者端与医院端共用服务端接口和数据库，患者预约成功后，医院端挂号管理与候诊队列能够刷新查看对应记录。由于本项目主要目标是桌面端课程设计，因此患者端不是独立手机App，而是作为患者预约入口集成在Qt客户端中。")
    add_figure(doc, images["patient_ui"], "图4-6 患者端预约界面布局图", 15.2)
    add_paragraph(doc, "患者端登录后只能访问个人相关功能，不能进入医院端管理模块。预约挂号时，患者端需要选择科室、医生、日期和号源，服务端会校验患者账号与患者档案的绑定关系，防止一个患者账号操作不属于自己的患者档案。预约成功后生成挂号记录和账单，医院端刷新后即可看到该挂号信息。")
    add_code_listing(doc, "文件4-6 患者端会话接口片段", """
class ApiClient : public QObject
{
public:
    void setPatientSession(const QJsonObject& data);
    bool isPatientLoggedIn() const;
    bool send(const common::Request& request);
private:
    QString m_patientToken;
    QJsonObject m_patientProfile;
};
""")
    add_paragraph(doc, "患者端会话与医院端工作人员会话分开保存，便于客户端判断当前入口身份。服务端收到患者端请求后，仍会根据Token确认患者账号身份，并在挂号服务中校验patient_id归属。")
    add_heading(doc, "4.5 本章小结", 2)
    add_paragraph(doc, "本章说明了系统开发环境、工程目录和核心模块实现。系统实现与第三章设计保持一致，通过Qt客户端、Qt服务端和MySQL数据库协同工作，完成了门诊挂号与药品管理的主要业务流程。")

    add_heading(doc, "第五章 结论", 1)
    add_paragraph(doc, "本课程设计完成了基于Qt和MySQL的医院门诊挂号与药品管理系统的需求分析、总体设计、数据库设计、模块实现和测试验证。系统采用客户端/服务端分离架构，客户端提供面向不同角色的桌面工作台和患者预约入口，服务端通过统一路由处理业务请求，数据库保存门诊核心数据。系统实现了患者档案、挂号排班、候诊接诊、病历检查、处方药品、收费退费、费用统计、权限配置和操作日志等功能，基本形成了从患者预约到诊疗收费的业务闭环。")
    add_paragraph(doc, "在实现过程中，项目重点解决了模块划分、角色权限、号源状态、处方流转、库存预警、支付状态和审计追溯等问题。通过CMake组织工程、跨平台脚本辅助构建、MySQL脚本初始化数据以及多项测试文件验证核心逻辑，系统具备较好的可维护性和演示完整性。")
    add_paragraph(doc, "由于课程设计时间和规模有限，系统仍有进一步完善空间。例如可以增加真实医保接口、短信通知、移动端小程序、电子签名、药品批次追踪、更多统计图表和更严格的并发压测。后续若继续扩展，可在现有模块化服务和数据库结构基础上逐步完善，使系统更接近实际医院门诊信息化平台。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Blanchette J, Summerfield M. C++ GUI Programming with Qt 4[M]. Prentice Hall, 2008.",
        "[2] MySQL 8.0 Reference Manual[EB/OL]. Oracle Corporation.",
        "[3] Qt Documentation. Qt Network and Widgets Modules[EB/OL]. The Qt Company.",
        "[4] 萨师煊, 王珊. 数据库系统概论[M]. 北京: 高等教育出版社.",
        "[5] Ian Sommerville. Software Engineering[M]. Pearson Education.",
        "[6] 国家卫生健康委员会. 医院信息化建设相关规范与实践资料[Z].",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False)

    add_heading(doc, "致谢", 1)
    add_paragraph(doc, "本课程设计从需求分析、系统设计到编码实现和文档整理，得到了指导教师和同学的帮助。在项目开发过程中，通过查阅Qt、MySQL和软件工程相关资料，进一步理解了客户端/服务端架构、数据库建模和业务模块划分方法。感谢老师在课程实践中的指导，也感谢同学在测试和演示过程中的建议。通过本次项目综合实践，我对医院门诊业务流程和信息系统开发有了更系统的认识，也提升了工程组织、问题分析和文档表达能力。")

    doc.save(ACTIVE_OUTPUT)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="生成课程设计总结 DOCX。")
    parser.add_argument("--template", default=str(TEMPLATE), help="课程设计总结模板 DOCX 路径。")
    parser.add_argument("--output", default=str(OUTPUT), help="主 DOCX 输出路径。")
    args = parser.parse_args()

    ACTIVE_TEMPLATE = Path(args.template)
    ACTIVE_OUTPUT = Path(args.output)
    stem = ACTIVE_OUTPUT.with_suffix("")
    ACTIVE_APPENDIX_OUTPUT = stem.with_name(stem.name + "-附件").with_suffix(".docx")
    ACTIVE_REFERENCES_OUTPUT = stem.with_name(stem.name + "-文献核验清单").with_suffix(".json")
    ACTIVE_SUMMARY_MD = stem.with_suffix(".md")
    ACTIVE_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document()
    print(ACTIVE_OUTPUT)
