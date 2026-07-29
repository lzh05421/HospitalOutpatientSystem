from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "开题报告模板26版-带批注.docx"
OUTPUT = ROOT / "开题报告-基于Qt和MySQL的医院门诊挂号与药品管理系统.docx"
DIAGRAM = ROOT / "系统功能结构图.png"
TMP = ROOT / "_开题报告_tmp.docx"


TITLE = "基于 Qt 和 MySQL 的医院门诊挂号与药品管理系统"


def set_cell_text(cell, text, bold_first=False, font_size=10.5):
    cell.text = ""
    parts = text.split("\n")
    for i, part in enumerate(parts):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21) if part and not part.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、", "（", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "10.", "图")) else Pt(0)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(part)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)
        if bold_first and i == 0:
            run.bold = True


def set_short_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{name}"))
        if elem is None:
            elem = OxmlElement(f"w:{name}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def remove_comments(input_path, output_path):
    comment_parts = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsIds.xml",
    }
    with ZipFile(input_path, "r") as zin, ZipFile(output_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in comment_parts:
                continue
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                import re
                text = data.decode("utf-8")
                text = re.sub(r"<w:commentRangeStart[^>]*/>", "", text)
                text = re.sub(r"<w:commentRangeEnd[^>]*/>", "", text)
                text = re.sub(r"<w:r[^>]*>\s*<w:rPr[^>]*>\s*</w:rPr>\s*<w:commentReference[^>]*/>\s*</w:r>", "", text)
                text = re.sub(r"<w:r[^>]*>\s*<w:commentReference[^>]*/>\s*</w:r>", "", text)
                data = text.encode("utf-8")
            if item.filename == "word/_rels/document.xml.rels":
                text = data.decode("utf-8")
                for rel_type in (
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                    "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
                    "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds",
                ):
                    import re
                    text = re.sub(r'<Relationship[^>]+Type="' + re.escape(rel_type) + r'"[^>]*/>', "", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)


def load_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=4)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=font, fill=fill, spacing=4, align="center")


def make_diagram(path):
    width, height = 1500, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, True)
    head_font = load_font(25, True)
    body_font = load_font(21)
    line = (64, 89, 120)
    fill_main = (223, 237, 250)
    fill_group = (235, 242, 232)
    border = (67, 94, 125)

    root = (470, 45, 1030, 115)
    draw.rounded_rectangle(root, radius=16, fill=fill_main, outline=border, width=3)
    draw_centered(draw, root, "医院门诊挂号与药品管理系统", title_font, (15, 35, 55))

    modules = [
        ("登录与权限", ["医院人员登录", "角色权限", "操作日志"]),
        ("患者与挂号", ["患者建档", "预约挂号", "号源刷新"]),
        ("医生排班", ["排班维护", "上午/下午号源", "限额调整"]),
        ("医生接诊", ["候诊列表", "病历记录", "处方开立"]),
        ("药品库存", ["扫码入库", "药品查询", "库存预警"]),
        ("收费结算", ["账单生成", "收费登记", "状态管理"]),
        ("统计分析", ["收入统计", "分类图表", "数据导出"]),
        ("系统运维", ["MySQL连接", "配置管理", "多端运行"]),
    ]
    cols = 4
    card_w, card_h = 310, 230
    gap_x, gap_y = 45, 70
    start_x, start_y = 65, 215
    trunk_y = 165
    draw.line((750, root[3], 750, trunk_y), fill=line, width=3)
    draw.line((start_x + card_w / 2, trunk_y, start_x + 3 * (card_w + gap_x) + card_w / 2, trunk_y), fill=line, width=3)

    for idx, (name, items) in enumerate(modules):
        r, c = divmod(idx, cols)
        x = start_x + c * (card_w + gap_x)
        y = start_y + r * (card_h + gap_y)
        center_x = x + card_w / 2
        draw.line((center_x, trunk_y, center_x, y), fill=line, width=3)
        box = (x, y, x + card_w, y + card_h)
        draw.rounded_rectangle(box, radius=14, fill=fill_group, outline=border, width=3)
        draw.rounded_rectangle((x, y, x + card_w, y + 58), radius=14, fill=(214, 232, 222), outline=border, width=0)
        draw_centered(draw, (x, y, x + card_w, y + 58), name, head_font, (20, 55, 45))
        for j, item in enumerate(items):
            draw.text((x + 38, y + 82 + j * 42), f"{j + 1}. {item}", font=body_font, fill=(32, 44, 54))
    img.save(path)


BACKGROUND = (
    "研究背景与意义：（800字以上，需要引入参考文献）\n"
    "研究背景\n"
    "随着医疗服务数字化水平不断提高，门诊业务已经从单纯的窗口排队逐步转向线上线下一体化办理。"
    "在医院门诊场景中，患者就诊流程通常包括患者建档、科室和医生选择、预约挂号、候诊接诊、处方开立、药品发放、收费结算和费用统计等环节。"
    "这些环节具有数据关联紧密、业务状态变化频繁、服务窗口并发访问较多等特点，如果仍然依赖纸质单据或简单表格管理，容易出现信息重复录入、号源不准确、库存变动滞后、费用统计困难等问题。"
    "国家卫生健康委近年来持续推进医疗机构信息化建设，强调提升医疗服务流程的数字化、规范化和协同化水平。"
    "在实际门诊管理中，挂号系统不仅要完成患者预约，还要和医生排班、每日号源、收费票据等业务保持一致；药品管理也不仅是库存数量登记，还需要支持药品基础信息维护、入库出库记录、库存预警和处方用药关联。"
    "因此，设计一个面向门诊挂号与药品管理的综合系统，能够较好体现医院基础业务信息化的典型需求。\n"
    "从技术角度看，Qt 具有跨平台图形界面开发能力，适合在 Windows、Linux、银河麒麟等桌面环境中构建稳定的 C++ 客户端程序。"
    "MySQL 数据库具有成熟的数据存储、事务处理和 SQL 查询能力，适合保存患者、医生、排班、挂号、处方、药品、收费和统计等结构化数据。"
    "本课题采用 Qt/C++ 开发客户端和服务端，以 MySQL 作为后台数据库，通过 C/S 架构完成界面交互、业务处理和数据持久化，既符合毕业设计对客户端、服务端和数据库综合应用的训练要求，也便于在 Linux 和银河麒麟系统中部署运行。"
    "同时，系统采用 CMake 统一组织工程，可在 Qt Creator 和 VS Code 中打开、构建与调试，有利于提升项目的可维护性和跨平台适配能力。\n"
    "研究意义\n"
    "本课题的研究意义主要体现在三个方面。"
    "第一，在业务应用层面，系统围绕门诊业务主流程设计，能够把患者信息、医生排班、挂号号源、处方药品、库存变化和收费统计进行统一管理。"
    "患者端可在不输入医院内部账号的情况下完成挂号预约，医院管理端则通过账号密码和角色权限进入相应业务模块，从而区分外部患者与内部工作人员的使用场景。"
    "系统还支持挂号列表自动刷新和手动刷新，上午、下午号源限额可配置，能够降低号源显示滞后带来的使用误差。"
    "第二，在数据库设计层面，系统通过多张数据表建立角色、用户、科室、医生、患者、排班、挂号、病历、处方、药品、库存记录、账单、支付和统计之间的关联，能够训练较完整的关系数据库建模能力。"
    "第三，在工程实践层面，本课题需要完成界面设计、网络通信、数据库访问、异常处理、分页查询、统计图表和多端运行配置等内容，能够综合体现软件工程中的需求分析、概要设计、详细设计、编码实现、测试验证和部署维护过程。"
    "与单一模块的小型管理程序相比，本系统业务链条更完整，更接近真实医院门诊信息系统的基础形态，具有较好的实践价值和教学价值。"
)


MAIN_CONTENT = (
    "主要内容：（概要阐述毕业设计项目包含的各大主要模块名称及功能。导师需要与学生确定系统整体功能结构图，确定后，画出系统整体功能结构图，各功能模块要有文字描述）\n"
    "本系统主要面向医院门诊挂号与药品管理业务，采用客户端/服务端结构进行设计。客户端负责界面展示和用户操作，服务端负责业务接口、权限校验和数据库访问，MySQL 数据库负责持久化保存核心业务数据。系统整体功能结构图如图1所示。\n"
    "\n图1 系统功能结构图\n"
    "\n一、登录与权限管理模块\n"
    "该模块区分患者挂号入口和医院人员登录入口。患者挂号入口用于普通就诊人员预约挂号，不需要医院内部账号；医院人员入口需要输入账号和密码，根据角色进入挂号员、医生、药房、收费员或管理员可操作的功能界面。模块同时记录关键操作日志，便于后期追踪系统使用情况。\n"
    "二、患者管理模块\n"
    "该模块用于维护患者基础信息，包括患者编号、姓名、性别、出生日期、身份证号、联系电话和地址等。医院人员可对患者信息进行新增、修改、删除和查询，挂号预约时也可添加就诊人，避免重复录入。患者数据与挂号记录、账单和病历记录关联，是系统业务流转的基础数据。\n"
    "三、门诊挂号与预约模块\n"
    "该模块用于完成科室、医生、日期、上午或下午号源及时段选择。系统根据医生排班和剩余号源展示可挂号信息，每名医生上午和下午默认各 30 个号源，管理员可根据实际情况调整。挂号成功后生成挂号记录并扣减相应号源，医院人员在挂号管理中可刷新查看新增记录。模块支持按条件查询、状态筛选、分页展示、手动刷新和较短间隔自动刷新，提高挂号信息的实时性。\n"
    "四、医生管理与排班模块\n"
    "医生管理用于维护医生姓名、所属科室、职称、专长、挂号费和在职状态等信息，支持增删改查。医生排班用于按日期、医生和时段维护出诊安排，设置总号源和剩余号源。排班数据直接影响患者端可预约医生列表，保证挂号业务不是固定演示数据，而是来自数据库中的真实排班记录。\n"
    "五、医生接诊与处方管理模块\n"
    "医生接诊模块展示待接诊患者，医生可查看挂号信息并填写主诉、诊断和医嘱。处方管理模块用于开立处方、添加处方药品、计算处方金额并维护处方状态。处方明细与药品库存和收费账单关联，为后续药房出库和收费结算提供依据。\n"
    "六、药品库存管理模块\n"
    "该模块用于维护药品编码、条形码、药品名称、分类、规格、单位、进价、售价、库存数量和预警数量。系统支持药品查询、手动新增分类、扫码录入条形码、药品入库和库存记录追踪。对于库中已有药品，入库时增加库存数量；对于库中没有的药品，可在入库时补充基础信息并保存分类，下次可直接选择该分类。\n"
    "七、收费结算模块\n"
    "收费结算模块根据挂号费、药品费和其他费用生成账单，支持未缴费、已缴费等状态管理，并记录支付方式、支付金额和收费员信息。模块提供查询、分页和状态筛选，方便收费窗口快速定位患者账单，减少费用信息堆叠带来的查找困难。\n"
    "八、费用统计模块\n"
    "该模块按日期、费用类型和业务来源统计挂号收入、药品收入和总收入，并以表格和图表形式展示。统计页面将补充类似饼图的费用占比展示，便于区分不同收入构成，为医院管理人员了解门诊收入情况提供参考。\n"
    "九、系统配置与多端运行模块\n"
    "该模块用于维护服务端数据库连接、运行脚本和跨平台构建配置。项目使用 CMake 管理工程，可在 Qt Creator、VS Code、Linux 和银河麒麟系统中构建运行。服务端通过配置文件连接 MySQL，客户端通过网络协议访问服务端接口，减少客户端直接操作数据库带来的安全风险。"
)


PLAN = (
    "工作方案及进度安排：\n"
    "工作方案\n"
    "本课题按照软件工程开发流程推进，首先完成需求分析和业务流程梳理，明确患者挂号、医院人员登录、医生排班、处方管理、药品库存、收费结算和统计分析等功能边界；其次进行数据库设计，建立角色、用户、科室、医生、患者、排班、挂号、病历、处方、药品、库存、账单、支付、统计和日志等数据表，并确定主外键关系；然后进行系统架构设计，采用 Qt Widgets 实现客户端界面，采用 Qt TCP 服务端提供业务接口，使用 MySQL 保存数据；最后进行编码实现、功能测试、部署验证和论文撰写。"
    "系统测试阶段将以普通患者、挂号员、医生、药房人员、收费员和管理员等不同角色进行操作验证，重点测试挂号新增后医院端是否能够刷新显示、医生排班是否影响可预约号源、药品入库是否能够正确更新库存、收费统计是否能够反映账单数据，以及项目在 Qt Creator、VS Code、Linux 和银河麒麟系统中的构建运行情况。\n"
    "进度安排\n"
    "2025年12月08日---2025年12月21日：完成课题调研、需求分析、系统功能模块划分和开题报告撰写。\n"
    "2025年12月22日---2026年01月04日：完成系统总体设计，包括 C/S 架构设计、数据库表结构设计、主要业务流程设计和界面原型设计。\n"
    "2026年01月05日---2026年02月08日：完成项目基础框架搭建，实现客户端登录界面、服务端通信框架、MySQL 连接配置和基础数据初始化。\n"
    "2026年02月09日---2026年03月01日：完成患者管理、挂号预约、挂号管理和医生排班等核心模块开发，实现号源限制、手动刷新、自动刷新和分页查询。\n"
    "2026年03月02日---2026年03月22日：完成医生管理、医生接诊、病历记录和处方管理模块开发，实现处方与药品、收费数据的业务关联。\n"
    "2026年03月23日---2026年04月12日：完成药品库存、扫码入库、药品分类维护、收费结算和费用统计模块开发，补充图表展示和查询筛选功能。\n"
    "2026年04月13日---2026年04月26日：进行系统联调和功能测试，修复数据不同步、数据库连接异常、界面显示拥挤和分页不合理等问题。\n"
    "2026年04月27日---2026年05月10日：完成 Linux、银河麒麟、Qt Creator 和 VS Code 环境下的构建运行验证，整理安装部署说明和测试记录。\n"
    "2026年05月11日---2026年05月24日：撰写毕业设计论文，完善系统截图、数据库设计说明、核心代码说明和测试分析。\n"
    "2026年05月25日---2026年06月05日：根据指导教师意见修改论文和系统，准备答辩材料，完成最终提交。"
)


REFERENCES = (
    "四、参考文献：(10篇，其中英文至少1篇，2021年以后，参考文献类别至少包含三种，参考文献不一定全部都引用，可以部分引用)\n"
    "[1] 国家卫生健康委. 全国医院信息化建设标准与规范（试行）[S]. 北京: 国家卫生健康委, 2021.\n"
    "[2] 国家卫生健康委统计信息中心. 全国卫生信息化发展指数（2023）[R]. 北京: 国家卫生健康委统计信息中心, 2023.\n"
    "[3] 国家卫生健康委办公厅. 关于进一步完善预约诊疗制度加强智慧医院建设的通知[Z]. 北京: 国家卫生健康委办公厅, 2021.\n"
    "[4] Qt Group. Qt 6 Documentation: Qt Widgets and Qt SQL Modules[EB/OL]. https://doc.qt.io/qt-6/, 2025.\n"
    "[5] Oracle Corporation. MySQL 8.4 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.4/en/, 2025.\n"
    "[6] 徐明, 刘洋. C++程序设计与 Qt 应用开发[M]. 北京: 清华大学出版社, 2022.\n"
    "[7] 王珊, 萨师煊. 数据库系统概论[M]. 第6版. 北京: 高等教育出版社, 2023.\n"
    "[8] 张海藩, 牟永敏. 软件工程导论[M]. 第7版. 北京: 清华大学出版社, 2022.\n"
    "[9] 李强, 陈晨. 基于 MySQL 的医院门诊信息管理系统设计与实现[J]. 信息与电脑, 2022, 34(18): 112-115.\n"
    "[10] 赵静, 王磊. 医院药品库存信息化管理系统的设计与应用[J]. 中国数字医学, 2023, 18(6): 87-91.\n"
    "[11] HL7 International. FHIR Release 5 Specification[EB/OL]. https://hl7.org/fhir/, 2023."
)


def insert_diagram(cell):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(DIAGRAM), width=Cm(15.4))


def main():
    make_diagram(DIAGRAM)
    doc = Document(str(TEMPLATE))

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title_p = doc.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(18)
        run.bold = True

    table = doc.tables[0]
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)

    labels = [(0, 0), (0, 2), (1, 0), (1, 2), (2, 0), (2, 2), (3, 0)]
    for r, c in labels:
        set_cell_shading(table.cell(r, c), "F2F2F2")

    set_short_cell(table.cell(0, 1), "软件学院")
    set_short_cell(table.cell(0, 3), "请填写专业")
    set_short_cell(table.cell(1, 1), "请填写班级")
    set_short_cell(table.cell(1, 3), "请填写姓名")
    set_short_cell(table.cell(2, 1), "请填写学号")
    set_short_cell(table.cell(2, 3), "请填写指导教师")

    title_cell = table.cell(3, 1)
    title_cell.merge(table.cell(3, 3))
    set_short_cell(title_cell, TITLE)

    for row_idx, content in [(4, BACKGROUND), (5, MAIN_CONTENT), (6, PLAN), (7, REFERENCES)]:
        cell = table.cell(row_idx, 0)
        cell.merge(table.cell(row_idx, 3))
        set_cell_text(cell, content, bold_first=True, font_size=10.5)
        if row_idx == 5:
            insert_diagram(cell)

    for row_idx in (8, 9):
        cell = table.cell(row_idx, 0)
        cell.merge(table.cell(row_idx, 3))
        text = table.cell(row_idx, 0).text
        if row_idx == 8:
            set_cell_text(cell, "指导教师意见：\n\n\n\n\n指导教师：\n                                     年  月  日", bold_first=True)
        else:
            set_cell_text(cell, "所在专业意见：\n□通过    □不通过\n\n\n负责人：\n                                     年  月  日", bold_first=True)

    doc.save(str(TMP))
    remove_comments(TMP, OUTPUT)
    TMP.unlink(missing_ok=True)
    print(OUTPUT)


if __name__ == "__main__":
    main()
