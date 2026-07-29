from __future__ import annotations

import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"E:\15751\Downloads\论文模板260205 (1).docx")
STANDARD = Path(r"E:\15751\Downloads\论文撰写规范.docx")
OUTPUT = ROOT / "docs" / "医院门诊挂号与药品管理系统-论文初稿.docx"
WORKING = ROOT / "docs" / "医院门诊挂号与药品管理系统-论文初稿-working.docx"

TITLE = "基于Qt与MySQL的医院门诊挂号与药品管理系统的设计与实现"


def set_run_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line=True, line=23, before=0, after=0):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(line)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def clear_document_body(doc: Document):
    body = doc._element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def configure_doc(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.gutter = Cm(0)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)


def page_break(doc: Document):
    doc.add_page_break()


def add_center(doc: Document, text: str, size=12, bold=False, spacing_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = Pt(23)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_para(doc: Document, text: str = "", first_line=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if level <= 3 else "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(23)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=15, bold=True, font="黑体")
    elif level == 2:
        set_run_font(r, size=14, bold=True, font="黑体")
    else:
        set_run_font(r, size=12, bold=True, font="黑体")
    return p


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(18)
    r = p.add_run(str(text))
    set_run_font(r, size=10.5, bold=bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=10.5, font="宋体")
    return p


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_text(hdr[i], text, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    set_table_borders(table)
    return table


def add_toc_placeholder(doc: Document):
    add_heading(doc, "目录", 1)
    toc_items = [
        "摘要",
        "ABSTRACT",
        "第1章 绪论",
        "第2章 需求分析",
        "第3章 系统总体设计",
        "第4章 系统详细设计与实现",
        "第5章 系统测试",
        "第6章 总结与展望",
        "参考文献",
        "致谢",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(23)
        p.paragraph_format.first_line_indent = None
        r1 = p.add_run(item)
        r2 = p.add_run("\t")
        r3 = p.add_run("页码待更新")
        for r in (r1, r2, r3):
            set_run_font(r)
    add_para(doc, "说明：打开 Word 后可在“引用”选项卡中更新目录页码。", first_line=False)


def add_cover(doc: Document):
    add_center(doc, "毕  业  设  计", size=22, bold=True, spacing_after=36)
    add_center(doc, TITLE, size=18, bold=True, spacing_after=42)
    for label in [
        "学    院：________________",
        "专    业：________________",
        "学生姓名：________________",
        "学生学号：________________",
        "指导教师：________________",
    ]:
        p = add_center(doc, label, size=14, spacing_after=10)
        p.paragraph_format.line_spacing = Pt(28)
    add_center(doc, "二〇二六年五月", size=14, spacing_after=0)
    page_break(doc)

    add_heading(doc, "原创性声明", 1)
    add_para(doc, "本人郑重声明：所呈交的毕业设计论文，是本人在指导教师指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本文不包含任何其他个人或集体已经发表或撰写过的作品成果。对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。")
    add_para(doc, "作者签名：______________        日期：______年____月____日", first_line=False)
    add_heading(doc, "版权使用授权声明", 1)
    add_para(doc, "本人完全了解学校有关保留、使用毕业设计论文的规定，同意学校保留并向有关部门或机构送交论文的复印件和电子版，允许论文被查阅和借阅。本人授权学校可以将本论文的全部或部分内容编入有关数据库进行检索，可以采用影印、缩印或扫描等复制手段保存和汇编本论文。")
    add_para(doc, "作者签名：______________        指导教师签名：______________", first_line=False)
    page_break(doc)


def add_abstracts(doc: Document):
    add_heading(doc, "摘要", 1)
    add_para(doc, "随着医院门诊业务量持续增长，传统依赖人工登记、纸质单据和分散统计的门诊管理方式在挂号排队、医生排班、处方流转、药品库存和收费核对等环节容易出现效率不足、信息滞后和数据一致性差等问题。为提高门诊业务处理效率和基础数据管理水平，本文设计并实现了一套基于 Qt 与 MySQL 的医院门诊挂号与药品管理系统。系统采用客户端与服务端分离的架构，客户端基于 Qt Widgets 构建桌面端交互界面，服务端基于 Qt TCP 网络通信提供统一业务接口，客户端与服务端之间使用 JSON Lines 格式传输业务请求与响应，数据库采用 MySQL 存储患者、科室、医生、排班、挂号、病历、检查、处方、药品、库存、账单、支付、权限和操作日志等业务数据。")
    add_para(doc, "本文围绕门诊主流程完成需求分析、总体设计、数据库设计、核心模块实现与测试验证。系统实现了患者预约挂号、医院人员登录、患者档案维护、科室医生管理、医生排班、候诊队列、医生接诊、检查申请、处方审核发药、药品库存预警、收费结算、退费处理、费用统计、权限管理和操作日志等功能。针对门诊场景中重复提交、库存扣减、医保挂号校验、支付状态流转和审计追溯等问题，系统在服务端业务层增加了状态校验、操作日志、业务编号和事务化处理设计。测试结果表明，系统能够完成门诊挂号与药品管理的主要业务流程，界面操作清晰，数据结构较完整，能够满足毕业设计范围内医院门诊信息化管理的基本要求。")
    add_para(doc, "关键词：医院门诊；挂号管理；药品管理；Qt；MySQL；客户端服务端架构", first_line=False)
    page_break(doc)

    add_heading(doc, "ABSTRACT", 1)
    add_para(doc, "With the continuous growth of outpatient services, traditional manual registration, paper-based documents and scattered statistics can easily lead to low efficiency, delayed information and weak data consistency in appointment registration, doctor scheduling, prescription circulation, drug inventory and payment verification. To improve the efficiency of outpatient business processing and basic data management, this thesis designs and implements a hospital outpatient registration and drug management system based on Qt and MySQL. The system adopts a client-server architecture. The client is built with Qt Widgets, while the server provides unified business interfaces through Qt TCP communication. Business requests and responses are transmitted in JSON Lines format, and MySQL is used to store data such as patients, departments, doctors, schedules, registrations, medical records, examinations, prescriptions, drugs, inventory records, bills, payments, permissions and operation logs.")
    add_para(doc, "This thesis completes requirement analysis, overall design, database design, key module implementation and testing around the main outpatient workflow. The system supports patient appointment registration, staff login, patient record management, department and doctor management, doctor scheduling, waiting queue management, consultation, examination request, prescription review and dispensing, drug inventory warning, billing, refund processing, fee statistics, permission management and operation logging. For duplicate submissions, inventory deduction, insurance verification, payment status flow and audit traceability, the server-side business layer introduces status checks, operation logs, business numbers and transactional processing. The test results show that the system can complete the main outpatient registration and drug management processes, provide clear user interaction and maintain relatively complete data structures, meeting the basic requirements of outpatient information management within the scope of this graduation project.")
    add_para(doc, "Key words: outpatient service; registration management; drug management; Qt; MySQL; client-server architecture", first_line=False)
    page_break(doc)


def add_chapter_1(doc: Document):
    add_heading(doc, "第1章 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    for text in [
        "医院门诊是医疗机构中患者接触最频繁、业务流转最密集的场景之一。患者从建档、挂号、候诊、接诊、检查、处方、收费到取药，需要多个岗位协同完成。若各环节主要依赖人工记录或分散系统处理，容易出现排班号源不清晰、候诊状态更新不及时、处方与库存衔接不紧密、收费信息重复核对等问题，既影响患者就诊体验，也增加工作人员的业务压力。",
        "随着桌面应用开发技术、关系型数据库和网络通信技术的发展，中小型医疗信息系统可以通过客户端服务端架构实现业务集中处理和数据统一存储。Qt 具有跨平台、组件完整、界面开发效率较高等特点，适合构建医院工作站类桌面应用；MySQL 具有成熟稳定、维护成本较低和数据查询能力较强等特点，适合存储门诊业务数据。将 Qt 与 MySQL 结合，可以在毕业设计规模内完成一套结构清晰、可运行、可扩展的门诊业务系统。",
    ]:
        add_para(doc, text)
    add_heading(doc, "1.2 研究意义", 2)
    for text in [
        "本文研究的意义主要体现在三个方面。第一，系统围绕门诊实际业务流程进行设计，将挂号、排班、接诊、处方、库存和收费等环节串联起来，有助于理解医院门诊信息化系统的基本业务边界。第二，系统采用客户端与服务端分离的方式，将界面展示、业务处理和数据持久化进行分层，有助于提升系统可维护性。第三，系统在数据库设计中覆盖患者、医生、药品、账单、权限和日志等核心实体，能够为后续功能扩展和数据统计提供基础。",
        "对于毕业设计实践而言，本系统不仅关注页面数量，还关注业务状态、数据约束、测试验证和部署运行。通过该系统的设计与实现，可以综合运用 C++、Qt、MySQL、CMake、网络通信、数据库建模和软件测试等知识，体现软件工程方法在实际业务系统开发中的应用。",
    ]:
        add_para(doc, text)
    add_heading(doc, "1.3 国内外研究现状", 2)
    for text in [
        "国外医院信息系统发展较早，电子病历、预约挂号、药品管理、医保支付和临床决策支持等功能已逐步形成较完善的体系。大型医疗机构通常采用集成化医院信息平台，将 HIS、EMR、LIS、PACS 和药房管理等系统连接起来，实现跨部门的信息共享。相关系统强调数据标准化、互操作能力、权限控制和医疗质量追踪。",
        "国内医院信息化建设近年来发展迅速，线上预约、分时段就诊、电子处方、移动支付和医保结算等功能逐渐普及。但对于教学实践和中小型业务场景而言，仍需要一类轻量化、结构清晰、便于演示和维护的门诊管理系统。本文所设计的系统不追求覆盖大型医院全部业务，而是选择门诊挂号与药品管理两个核心方向，结合患者服务、医生工作站、药房和收费窗口等角色完成闭环设计。",
    ]:
        add_para(doc, text)
    add_heading(doc, "1.4 论文主要工作", 2)
    for text in [
        "本文主要完成以下工作：第一，分析医院门诊挂号与药品管理系统的业务需求，明确患者、挂号员、医生、药师、收费员和管理员等角色的功能范围；第二，设计系统总体架构，确定 Qt 客户端、Qt TCP 服务端、公共协议层和 MySQL 数据库的分层结构；第三，完成数据库概念结构与逻辑结构设计，设计患者、排班、挂号、病历、处方、药品、库存、账单、支付、权限和日志等数据表；第四，实现挂号管理、医生接诊、处方管理、药品库存和收费结算等核心模块；第五，通过单元测试、源码规则测试和业务流程测试验证系统主要功能。",
    ]:
        add_para(doc, text)
    add_heading(doc, "1.5 论文组织结构", 2)
    add_para(doc, "本文共分为六章。第1章介绍研究背景、研究意义、研究现状和论文主要工作；第2章进行系统需求分析；第3章介绍系统总体设计和数据库设计；第4章阐述核心功能模块的详细设计与实现；第5章进行系统测试；第6章总结全文工作并展望后续改进方向。")


def add_chapter_2(doc: Document):
    add_heading(doc, "第2章 需求分析", 1)
    add_heading(doc, "2.1 可行性分析", 2)
    add_para(doc, "技术可行性方面，系统采用 Qt/C++ 进行客户端和服务端开发，使用 CMake 组织工程，使用 MySQL 进行数据持久化。Qt 提供界面组件、网络通信、JSON 处理和数据库访问能力，能够支撑本系统的主要功能开发。经济可行性方面，系统所采用的开发工具和数据库均可用于教学和毕业设计环境，部署成本较低。操作可行性方面，系统采用桌面工作台界面，将模块导航、表格查询、表单录入和操作按钮集中呈现，符合医院窗口人员和医生工作站的常见使用方式。")
    add_heading(doc, "2.2 用户角色分析", 2)
    add_table(doc, "表2-1 系统用户角色及职责", ["角色", "主要职责", "典型功能"], [
        ["患者", "在线选择科室、医生和号源，完成预约挂号和就诊人管理", "患者登录、预约挂号、取消预约、就诊人维护"],
        ["挂号员", "维护患者档案，办理现场挂号和退号", "患者管理、挂号管理、候诊查询"],
        ["医生", "查看候诊队列，完成接诊、病历记录、检查申请和处方开立", "医生接诊、病历编辑、检查申请、处方创建"],
        ["药师", "审核处方，完成发药和退药，维护药品库存", "处方审核、发药、库存入库、库存预警"],
        ["收费员", "处理挂号费、检查费和药品费，完成支付和退费审核", "账单收费、支付状态查询、退费处理"],
        ["管理员", "维护科室、医生、角色权限和系统日志", "科室管理、医生管理、权限配置、操作日志"],
    ])
    add_heading(doc, "2.3 功能需求分析", 2)
    add_para(doc, "根据门诊业务流程，本系统划分为入口登录、患者管理、挂号管理、医生排班、候诊队列、医生接诊、检查管理、处方管理、药品库存、收费结算、费用统计、权限管理和操作日志等功能模块。各模块之间围绕挂号记录、处方记录、账单记录和库存记录进行数据关联。")
    add_table(doc, "表2-2 系统功能需求表", ["模块", "功能需求", "说明"], [
        ["登录权限", "医院人员登录、患者登录、会话保存、角色权限控制", "不同角色进入系统后展示与权限匹配的功能"],
        ["患者管理", "新增、查询、修改患者档案，维护就诊人信息", "支持患者基础资料和患者账号关联"],
        ["挂号管理", "科室选择、医生号源查询、预约挂号、现场挂号、退号", "挂号成功后生成挂号单号并扣减号源"],
        ["医生排班", "维护医生班次、号源数量、停诊规则和智能轮排", "支持按医生、科室和日期管理号源"],
        ["医生接诊", "候诊队列、叫号、接诊、病历记录、检查申请", "医生可围绕挂号记录完成诊疗操作"],
        ["处方管理", "处方开立、明细维护、处方审核、发药退药", "处方与药品库存和账单费用关联"],
        ["药品库存", "药品信息维护、入库、出库、库存预警", "库存变动写入库存流水"],
        ["收费结算", "账单生成、扫码支付、医保支付模拟、退费审核", "记录支付流水并更新账单状态"],
        ["统计日志", "费用统计、驾驶舱、操作日志和审计明细", "支持日常管理和问题追溯"],
    ])
    add_heading(doc, "2.4 非功能需求分析", 2)
    for text in [
        "可靠性需求：系统应在数据库连接异常时给出明确提示，并在服务端统一处理业务失败响应。挂号、支付、处方发药和库存扣减等关键操作应避免重复提交造成数据异常。",
        "易用性需求：客户端页面应保持统一的表格、按钮和表单风格，常用功能支持刷新、筛选、分页和复制关键编号，减少窗口人员重复操作。",
        "安全性需求：系统应通过账号密码和角色权限控制限制功能访问，关键操作写入操作日志，便于后续追溯。患者密码和业务支付令牌不应以明文方式在数据库中保存。",
        "可维护性需求：系统应按照客户端、公共协议、服务端业务模块和数据库脚本分层组织代码，新增模块时能够复用请求路由、表格页面和数据库访问模式。",
    ]:
        add_para(doc, text)
    add_heading(doc, "2.5 核心用例描述", 2)
    add_table(doc, "表2-3 患者预约挂号用例描述", ["项目", "内容"], [
        ["用例编号", "UC-REG-01"],
        ["用例名称", "预约挂号"],
        ["参与者", "患者、挂号员"],
        ["前置条件", "患者已登录或已完成患者档案登记，医生排班存在可用号源"],
        ["基本流程", "选择科室；选择医生和日期；查询可用号源；提交挂号信息；系统校验医保或自费信息；生成挂号记录和账单"],
        ["后置条件", "挂号状态为待就诊，号源余量减少，系统记录操作日志"],
    ])
    add_table(doc, "表2-4 医生接诊用例描述", ["项目", "内容"], [
        ["用例编号", "UC-CON-01"],
        ["用例名称", "接诊并填写病历"],
        ["参与者", "医生"],
        ["前置条件", "医生已登录，候诊队列存在待接诊患者"],
        ["基本流程", "查看候诊队列；选择患者；叫号或开始接诊；录入主诉、现病史、诊断和医嘱；根据病情开立检查或处方；保存病历"],
        ["后置条件", "挂号状态更新，病历记录保存，相关检查或处方进入后续处理流程"],
    ])
    add_table(doc, "表2-5 处方审核发药用例描述", ["项目", "内容"], [
        ["用例编号", "UC-RX-01"],
        ["用例名称", "审核处方并发药"],
        ["参与者", "药师"],
        ["前置条件", "医生已开立处方，处方明细中的药品库存满足发药要求"],
        ["基本流程", "药师查看待审核处方；核对药品、剂量和用法；审核通过或驳回；患者缴费后执行发药；系统扣减库存并写入库存流水"],
        ["后置条件", "处方状态更新为已发药，药品库存减少，操作日志记录审核和发药动作"],
    ])
    add_table(doc, "表2-6 收费结算用例描述", ["项目", "内容"], [
        ["用例编号", "UC-BIL-01"],
        ["用例名称", "账单收费"],
        ["参与者", "收费员"],
        ["前置条件", "挂号、检查或处方已形成待支付账单"],
        ["基本流程", "查询患者账单；核对费用明细；选择现金、微信、支付宝或医保模拟支付；系统记录支付流水；账单状态更新为已缴费"],
        ["后置条件", "账单支付完成，支付记录生成，处方或检查可进入后续环节"],
    ])
    add_table(doc, "表2-7 药品入库用例描述", ["项目", "内容"], [
        ["用例编号", "UC-INV-01"],
        ["用例名称", "药品入库"],
        ["参与者", "药师、管理员"],
        ["前置条件", "药品基础资料存在或已录入"],
        ["基本流程", "录入药品条码、规格、批次或数量；系统校验数量合法性；增加库存数量；生成入库流水；库存低于预警值时在库存列表提示"],
        ["后置条件", "药品库存更新，库存流水可查询"],
    ])


def add_chapter_3(doc: Document):
    add_heading(doc, "第3章 系统总体设计", 1)
    add_heading(doc, "3.1 系统设计目标", 2)
    add_para(doc, "系统设计目标是构建一套能够覆盖门诊挂号与药品管理主要流程的桌面应用系统。系统应具备清晰的模块划分、较完整的数据结构、可运行的客户端服务端通信能力和可验证的测试用例。系统重点解决患者挂号、医生排班、候诊接诊、处方流转、药品库存和收费统计之间的数据衔接问题。")
    add_heading(doc, "3.2 系统架构设计", 2)
    add_para(doc, "系统采用客户端服务端架构。客户端负责用户交互、表格展示、表单录入和请求发起；服务端负责业务路由、权限校验、数据访问和业务状态控制；公共协议层定义请求与响应结构；数据库层使用 MySQL 存储业务数据。客户端与服务端之间通过 TCP 长连接传输 JSON Lines 消息，每个请求包含模块名、动作名、请求编号和业务数据，服务端返回状态码、提示信息和响应数据。")
    add_caption(doc, "图3-1 系统总体架构图")
    add_para(doc, "客户端（Qt Widgets）  →  公共协议层（JSON Lines）  →  服务端（Qt TCP、RequestRouter、ModuleService）  →  数据库（MySQL）", first_line=False)
    add_heading(doc, "3.3 功能结构设计", 2)
    add_para(doc, "系统功能结构以门诊业务流程为主线，辅以基础数据、权限日志和统计分析模块。患者端主要完成预约挂号和就诊人维护；医院端围绕挂号员、医生、药师、收费员和管理员提供对应工作台。")
    add_table(doc, "表3-1 系统模块结构", ["一级模块", "二级功能"], [
        ["入口与登录", "患者入口、医院人员入口、角色预选、账号密码登录、会话保存"],
        ["门诊主流程", "预约挂号、挂号管理、候诊队列、医生接诊、检查管理、处方管理、收费结算"],
        ["基础数据", "患者管理、患者病历档案、科室管理、医生管理、医生排班"],
        ["药品与收费", "药品库存、处方审核、发药退药、账单收费、退费、费用统计"],
        ["管理支撑", "院长驾驶舱、权限控制、操作日志、审计明细、通用筛选、分页、导出"],
        ["运维支撑", "CMake 工程、Windows/Linux/银河麒麟脚本、MySQL 配置、演示模式"],
    ])
    add_heading(doc, "3.4 数据库概念结构设计", 2)
    add_para(doc, "数据库概念结构围绕患者、医生、科室、排班、挂号、病历、检查、处方、药品、库存、账单、支付、用户角色和日志等实体展开。患者与挂号记录是一对多关系，医生与排班记录是一对多关系，挂号记录与病历记录是一对一关系，挂号记录与处方、检查和账单存在关联关系，处方与处方明细是一对多关系，药品与处方明细、库存流水存在关联关系。")
    add_caption(doc, "图3-2 核心实体关系图")
    add_para(doc, "患者 -- 挂号 -- 医生/排班 -- 病历/检查/处方 -- 账单/支付；处方 -- 处方明细 -- 药品 -- 库存流水。", first_line=False)
    add_heading(doc, "3.5 数据库逻辑结构设计", 2)
    add_para(doc, "系统数据库名为 hospital_outpatient。根据当前项目数据库脚本，系统核心业务表覆盖角色用户、组织权限、患者账号、挂号医保校验、医生排班、病历、检查、处方、药品、库存、账单、支付、费用统计、操作日志和事务事件等内容。")
    add_table(doc, "表3-2 患者表 patients", ["字段名", "类型", "长度", "主键", "允许为空", "说明"], [
        ["id", "BIGINT", "-", "是", "否", "患者主键"],
        ["patient_no", "VARCHAR", "32", "否", "否", "患者编号，唯一"],
        ["user_id", "BIGINT", "-", "否", "是", "关联患者账号"],
        ["name", "VARCHAR", "64", "否", "否", "患者姓名"],
        ["gender", "VARCHAR", "8", "否", "否", "性别"],
        ["birthday", "DATE", "-", "否", "是", "出生日期"],
        ["id_card", "VARCHAR", "32", "否", "是", "身份证号"],
        ["phone", "VARCHAR", "32", "否", "是", "联系电话"],
        ["relationship", "VARCHAR", "32", "否", "是", "与账号持有人关系"],
        ["address", "VARCHAR", "255", "否", "是", "联系地址"],
    ])
    add_table(doc, "表3-3 医生排班表 doctor_schedules", ["字段名", "类型", "长度", "主键", "允许为空", "说明"], [
        ["id", "BIGINT", "-", "是", "否", "排班主键"],
        ["doctor_id", "BIGINT", "-", "否", "否", "医生编号"],
        ["department_id", "BIGINT", "-", "否", "是", "科室编号"],
        ["work_date", "DATE", "-", "否", "否", "出诊日期"],
        ["period", "VARCHAR", "16", "否", "否", "出诊时段"],
        ["total_quota", "INT", "-", "否", "否", "总号源数"],
        ["remain_quota", "INT", "-", "否", "否", "剩余号源数"],
        ["status", "TINYINT", "-", "否", "否", "排班状态"],
    ])
    add_table(doc, "表3-4 挂号表 registrations", ["字段名", "类型", "长度", "主键", "允许为空", "说明"], [
        ["id", "BIGINT", "-", "是", "否", "挂号主键"],
        ["registration_no", "VARCHAR", "32", "否", "否", "挂号单号，唯一"],
        ["patient_id", "BIGINT", "-", "否", "否", "患者编号"],
        ["doctor_id", "BIGINT", "-", "否", "否", "医生编号"],
        ["schedule_id", "BIGINT", "-", "否", "否", "排班编号"],
        ["appointment_time_slot", "VARCHAR", "32", "否", "是", "预约时段"],
        ["register_time", "DATETIME", "-", "否", "否", "挂号时间"],
        ["status", "VARCHAR", "16", "否", "否", "挂号状态"],
        ["fee", "DECIMAL", "10,2", "否", "否", "挂号费用"],
        ["operator_id", "BIGINT", "-", "否", "否", "操作员编号"],
    ])
    add_table(doc, "表3-5 处方表 prescriptions", ["字段名", "类型", "长度", "主键", "允许为空", "说明"], [
        ["id", "BIGINT", "-", "是", "否", "处方主键"],
        ["prescription_no", "VARCHAR", "32", "否", "否", "处方编号，唯一"],
        ["registration_id", "BIGINT", "-", "否", "否", "关联挂号记录"],
        ["doctor_id", "BIGINT", "-", "否", "否", "开方医生"],
        ["status", "VARCHAR", "16", "否", "否", "处方状态"],
        ["total_amount", "DECIMAL", "10,2", "否", "否", "处方总金额"],
        ["reviewer_id", "BIGINT", "-", "否", "是", "审核药师"],
        ["dispense_user_id", "BIGINT", "-", "否", "是", "发药人员"],
        ["created_at", "DATETIME", "-", "否", "否", "创建时间"],
    ])
    add_table(doc, "表3-6 药品表 drugs", ["字段名", "类型", "长度", "主键", "允许为空", "说明"], [
        ["id", "BIGINT", "-", "是", "否", "药品主键"],
        ["drug_code", "VARCHAR", "32", "否", "否", "药品编号，唯一"],
        ["barcode", "VARCHAR", "64", "否", "是", "条码"],
        ["drug_name", "VARCHAR", "128", "否", "否", "药品名称"],
        ["category_id", "BIGINT", "-", "否", "否", "药品分类"],
        ["specification", "VARCHAR", "128", "否", "是", "规格"],
        ["unit", "VARCHAR", "16", "否", "否", "单位"],
        ["sale_price", "DECIMAL", "10,2", "否", "否", "销售价格"],
        ["stock_quantity", "INT", "-", "否", "否", "库存数量"],
        ["warning_quantity", "INT", "-", "否", "否", "预警数量"],
    ])
    add_table(doc, "表3-7 账单表 bills", ["字段名", "类型", "长度", "主键", "允许为空", "说明"], [
        ["id", "BIGINT", "-", "是", "否", "账单主键"],
        ["bill_no", "VARCHAR", "32", "否", "否", "账单编号，唯一"],
        ["registration_id", "BIGINT", "-", "否", "否", "挂号编号"],
        ["patient_id", "BIGINT", "-", "否", "否", "患者编号"],
        ["registration_fee", "DECIMAL", "10,2", "否", "否", "挂号费"],
        ["drug_fee", "DECIMAL", "10,2", "否", "否", "药品费"],
        ["other_fee", "DECIMAL", "10,2", "否", "否", "其他费用"],
        ["total_amount", "DECIMAL", "10,2", "否", "否", "总金额"],
        ["status", "VARCHAR", "32", "否", "否", "账单状态"],
        ["pay_time", "DATETIME", "-", "否", "是", "支付时间"],
    ])
    add_heading(doc, "3.6 通信协议设计", 2)
    add_para(doc, "系统公共协议层定义 Request 与 Response 结构。客户端请求包含 module、action、requestId 和 payload 字段，服务端响应包含 success、message 和 data 字段。module 用于定位业务模块，action 表示具体操作，payload 承载业务数据。服务端 RequestRouter 根据 module 将请求分发到对应 ModuleService，业务服务完成参数校验、数据库操作和响应组装。")
    add_table(doc, "表3-8 通信请求字段说明", ["字段", "类型", "说明"], [
        ["module", "字符串", "业务模块，例如 registration、prescription、billing"],
        ["action", "字符串", "业务动作，例如 list、create、update、pay"],
        ["requestId", "字符串", "请求编号，用于追踪请求"],
        ["payload", "JSON对象", "业务参数"],
    ])


def add_chapter_4(doc: Document):
    add_heading(doc, "第4章 系统详细设计与实现", 1)
    add_heading(doc, "4.1 开发环境与工程结构", 2)
    add_para(doc, "系统使用 CMake 组织工程，主要目录包括 client、server、common、database、docs、scripts 和 tests。client 目录存放 Qt Widgets 客户端页面和网络请求封装；server 目录存放 TCP 服务端、请求路由、业务服务和数据库访问代码；common 目录存放客户端与服务端共享的协议结构；database 目录存放 MySQL 建库脚本和迁移脚本；tests 目录存放协议、路由、业务规则和源码规则测试。")
    add_table(doc, "表4-1 工程目录说明", ["目录", "说明"], [
        ["client", "Qt Widgets 客户端，包含登录、主窗口、患者端预约和各业务页面"],
        ["server", "Qt TCP 服务端，包含配置、连接管理、路由、权限和业务模块"],
        ["common", "公共协议和通用类型"],
        ["database", "MySQL 建库脚本和迁移脚本"],
        ["config", "服务端配置模板"],
        ["scripts", "Windows、Linux 和银河麒麟构建运行脚本"],
        ["tests", "自动化测试和源码规则测试"],
    ])
    add_heading(doc, "4.2 登录与权限模块实现", 2)
    add_para(doc, "登录模块包括医院人员登录和患者登录两类入口。医院人员登录后，系统根据用户角色展示挂号员、医生、药师、收费员或管理员可访问的模块。服务端通过 AuthService 和 AuthorizationService 处理账号校验、会话状态和权限访问控制。权限相关数据包括 sys_user、sys_role、sys_menu、sys_user_role、sys_role_menu 和 sys_user_dept 等表，能够支持用户、角色、菜单和部门之间的关联。")
    add_caption(doc, "图4-1 登录界面截图")
    add_para(doc, "截图插入位：请在系统运行后截取医院人员登录或患者登录界面，并替换本段。", first_line=False)
    add_heading(doc, "4.3 患者与挂号模块实现", 2)
    add_para(doc, "患者与挂号模块是系统的起始业务环节。患者管理页面支持患者档案的新增、查询和修改；患者端预约页面支持选择就诊人、科室、医生、日期和号源时段。挂号请求提交后，服务端 RegistrationService 校验患者、医生、排班和号源信息，生成挂号单号并保存 registrations 记录，同时根据费用信息形成账单基础数据。系统还设计了 registration_insurance_check、registration_insurance_tokens 和 registration_insurance_audit_logs 等表，用于模拟医保挂号资格校验和审计记录。")
    add_caption(doc, "图4-2 预约挂号界面截图")
    add_para(doc, "截图插入位：请插入患者端预约挂号界面，注意截屏时间应符合毕业设计时间范围。", first_line=False)
    add_heading(doc, "4.4 医生排班模块实现", 2)
    add_para(doc, "医生排班模块通过 SchedulePage 提供排班维护界面，服务端 ScheduleService 负责排班查询、创建和规则处理。排班数据保存于 doctor_schedules 表，包含医生、科室、出诊日期、时段、总号源和剩余号源。系统还设计 schedule_rules 表，用于保存停诊、轮排或规则化排班条件。为避免同一医生同一日期同一时段重复排班，doctor_schedules 表设置 doctor_id、work_date 和 period 的唯一约束。")
    add_heading(doc, "4.5 医生接诊与病历模块实现", 2)
    add_para(doc, "医生接诊模块围绕候诊队列展开。医生登录后可以查看当日待接诊、已叫号、检查中和已完成等状态的挂号记录。开始接诊后，医生录入主诉、现病史、既往史、体格检查、诊断和医嘱，并可根据病情开立检查项目或处方。病历数据保存于 medical_records 表，检查申请保存于 examinations 表。该设计使病历、检查和处方均能关联到同一挂号记录，方便回溯一次门诊就诊过程。")
    add_caption(doc, "图4-3 医生接诊界面截图")
    add_para(doc, "截图插入位：请插入候诊队列或医生接诊页面截图。", first_line=False)
    add_heading(doc, "4.6 处方与药品库存模块实现", 2)
    add_para(doc, "处方模块包括处方创建、处方明细维护、处方审核、发药和退药等环节。处方主表 prescriptions 保存处方编号、挂号记录、医生、状态、总金额、审核人和发药人等信息；prescription_items 保存药品、数量、用法用量、单价和金额。药品库存模块通过 drugs、drug_categories 和 stock_records 等表管理药品基础资料、库存数量和库存变动流水。发药时系统根据处方明细扣减库存，并写入 stock_records，保证药品库存变化可追踪。")
    add_caption(doc, "图4-4 药品库存界面截图")
    add_para(doc, "截图插入位：请插入药品库存列表或处方审核发药页面截图。", first_line=False)
    add_heading(doc, "4.7 收费结算模块实现", 2)
    add_para(doc, "收费结算模块由 BillingService 提供服务端业务处理，支持账单查询、账单更新、自费支付、医保模拟支付、支付二维码创建、支付状态查询、医保回调、退费申请和退费审核等操作。账单主表 bills 保存挂号费、药品费、检查费和总金额，payments 表保存支付流水，insurance_transactions 表保存医保交易模拟记录。系统通过账单状态控制收费流程，避免重复支付、重复退费等异常操作。")
    add_caption(doc, "图4-5 收费结算界面截图")
    add_para(doc, "截图插入位：请插入收费结算或支付窗口截图。", first_line=False)
    add_heading(doc, "4.8 统计、日志与系统管理模块实现", 2)
    add_para(doc, "统计模块提供日收入、科室收入和药品销售等数据展示，为管理人员了解门诊运行情况提供依据。操作日志模块通过 operation_logs 和 audit_log_details 记录用户在挂号、排班、接诊、处方、收费等模块中的关键操作。权限管理模块通过用户、角色、菜单和部门数据实现功能权限控制。系统管理模块还包括科室管理、医生管理和基础字典维护等内容，为门诊业务提供基础数据支撑。")


def add_chapter_5(doc: Document):
    add_heading(doc, "第5章 系统测试", 1)
    add_heading(doc, "5.1 测试环境", 2)
    add_para(doc, "系统测试在 Windows 开发环境下进行，使用 Qt 6、MinGW、CMake、MySQL 和 Qt Test 等工具。项目同时提供 Windows、Linux 和银河麒麟环境的配置、构建与运行脚本。测试内容包括公共协议测试、认证路由测试、排班规则测试、收费路由测试、数据库兼容性测试以及多个页面和业务流程的源码规则测试。")
    add_table(doc, "表5-1 测试环境说明", ["项目", "内容"], [
        ["操作系统", "Windows 开发环境，兼容 Linux/银河麒麟部署脚本"],
        ["开发框架", "Qt Widgets、Qt Network、Qt SQL"],
        ["数据库", "MySQL，Qt QODBC 或 QMYSQL 驱动"],
        ["构建工具", "CMake、MinGW"],
        ["测试工具", "Qt Test、CTest、源码规则测试"],
    ])
    add_heading(doc, "5.2 测试方法", 2)
    add_para(doc, "本文采用黑盒测试与白盒测试相结合的方式。黑盒测试主要从用户角度验证登录、挂号、接诊、处方、发药、收费和统计等流程是否符合预期；白盒测试主要通过协议测试、服务端路由测试、排班规则测试和源码规则测试验证代码结构和关键逻辑。对于涉及数据库的功能，重点检查数据能否正确写入、状态能否正确流转、异常输入能否被拦截。")
    add_heading(doc, "5.3 功能测试用例", 2)
    add_table(doc, "表5-2 系统功能测试用例", ["编号", "测试模块", "测试内容", "预期结果", "结果"], [
        ["login_01", "登录权限", "输入正确管理员账号和密码登录系统", "登录成功，进入主工作台", "通过"],
        ["login_02", "登录权限", "输入错误密码登录系统", "登录失败并提示账号或密码错误", "通过"],
        ["gh_01", "挂号管理", "选择患者、科室、医生和可用号源提交挂号", "生成挂号单号，号源余量减少", "通过"],
        ["gh_02", "挂号管理", "对已取消或无号源排班提交挂号", "系统拒绝提交并显示错误提示", "通过"],
        ["pb_01", "医生排班", "新增医生当日排班和号源数量", "排班列表出现新记录", "通过"],
        ["jz_01", "医生接诊", "医生查看候诊队列并开始接诊", "挂号状态更新，病历编辑区域可用", "通过"],
        ["bl_01", "病历管理", "保存主诉、诊断和医嘱", "病历记录写入数据库并可再次查询", "通过"],
        ["cf_01", "处方管理", "医生开立包含多种药品的处方", "处方主表和明细表生成记录", "通过"],
        ["yf_01", "药房管理", "药师审核处方并执行发药", "处方状态更新，药品库存扣减", "通过"],
        ["kc_01", "药品库存", "录入药品入库数量", "库存数量增加并生成库存流水", "通过"],
        ["sf_01", "收费结算", "对待支付账单执行自费支付", "账单状态变为已缴费并生成支付记录", "通过"],
        ["tf_01", "收费结算", "对已缴费账单发起退费审核", "账单状态按退费流程流转", "通过"],
        ["tj_01", "费用统计", "查询当日费用统计", "返回挂号收入、药品收入和总收入", "通过"],
        ["qx_01", "权限管理", "普通角色访问未授权模块", "系统限制访问或不展示入口", "通过"],
    ])
    add_heading(doc, "5.4 自动化测试情况", 2)
    add_para(doc, "项目 tests 目录中包含多项自动化测试。protocol_tests 用于验证公共协议编解码；auth_router_tests 用于验证认证与路由处理；workflow_rules_tests 用于验证业务流程规则；schedule_rule_engine_tests 和 schedule_service_tests 用于验证排班规则和排班服务；payment_router_tests 用于验证支付相关路由；database_compatibility_schema_tests 用于验证数据库脚本兼容性；其余源码规则测试用于检查页面功能、支付安全、退费流程、药房流程、权限管理、候诊筛选和仪表盘可视化等内容。")
    add_table(doc, "表5-3 自动化测试项目示例", ["测试文件", "测试目标"], [
        ["protocol_tests.cpp", "验证请求响应协议结构"],
        ["auth_router_tests.cpp", "验证登录认证和请求路由"],
        ["schedule_rule_engine_tests.cpp", "验证排班规则计算"],
        ["schedule_batch_transaction_tests.cpp", "验证批量排班事务处理"],
        ["payment_router_tests.cpp", "验证支付业务路由"],
        ["registration_insurance_source_tests.cpp", "验证挂号医保校验相关源码规则"],
        ["pharmacy_workflow_source_tests.cpp", "验证药房处方审核发药流程"],
        ["database_compatibility_schema_tests.cpp", "验证数据库表结构兼容性"],
    ])
    add_heading(doc, "5.5 测试结果分析", 2)
    add_para(doc, "从功能测试结果看，系统能够完成从患者挂号、医生接诊、处方开立、药师发药到收费统计的主要业务流程。系统对无号源、错误登录、未授权访问、重复支付和库存不足等异常场景进行了提示或限制。自动化测试覆盖了协议、路由、排班、支付、数据库兼容性和若干源码规则，能够在后续修改代码时辅助发现回归问题。由于毕业设计环境与真实医院生产系统仍存在差距，系统在性能压测、真实医保接口、真实支付接口和大规模并发场景方面仍需进一步完善。")


def add_chapter_6(doc: Document):
    add_heading(doc, "第6章 总结与展望", 1)
    add_heading(doc, "6.1 总结", 2)
    add_para(doc, "本文围绕医院门诊挂号与药品管理业务，设计并实现了一套基于 Qt 与 MySQL 的门诊信息管理系统。系统采用客户端服务端架构，客户端使用 Qt Widgets 构建桌面工作台，服务端使用 Qt TCP 和模块化业务服务处理请求，数据库使用 MySQL 存储业务数据。系统实现了患者预约挂号、患者档案、科室医生、医生排班、候诊接诊、病历记录、检查申请、处方审核发药、药品库存、收费结算、费用统计、权限管理和操作日志等功能。")
    add_para(doc, "在设计过程中，本文完成了需求分析、功能结构设计、数据库设计、通信协议设计、核心模块实现和系统测试。数据库设计覆盖了门诊业务所需的主要实体和关系，服务端通过 RequestRouter 和 ModuleService 实现模块化分发，客户端通过统一页面和业务页面提升了功能复用性。测试结果表明，系统能够满足毕业设计范围内医院门诊挂号与药品管理的基本需求。")
    add_heading(doc, "6.2 展望", 2)
    add_para(doc, "后续可从以下方面继续完善系统。第一，引入更完整的电子病历结构化模板，提高病历数据的规范性和可检索性。第二，完善真实医保和第三方支付接口，替换当前模拟支付与医保校验流程。第三，增加消息提醒和实时推送能力，使候诊队列、处方审核和支付状态能够实时同步。第四，强化安全设计，对敏感信息进行更严格的脱敏、加密和访问审计。第五，进一步开展压力测试和多终端兼容测试，提高系统在真实业务场景中的稳定性和扩展能力。")


def add_references_and_ack(doc: Document):
    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] 张海藩. 软件工程导论[M]. 北京: 清华大学出版社, 2013.",
        "[2] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.",
        "[3] Mark Summerfield. Advanced Qt Programming[M]. Boston: Addison-Wesley, 2010.",
        "[4] Jasmin Blanchette, Mark Summerfield. C++ GUI Programming with Qt 4[M]. Upper Saddle River: Prentice Hall, 2008.",
        "[5] MySQL Documentation Team. MySQL 8.0 Reference Manual[EB/OL]. Oracle Corporation.",
        "[6] Qt Company. Qt Documentation[EB/OL]. The Qt Company.",
        "[7] Gamma E, Helm R, Johnson R, Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software[M]. Boston: Addison-Wesley, 1994.",
        "[8] Martin R C. Clean Code: A Handbook of Agile Software Craftsmanship[M]. Upper Saddle River: Prentice Hall, 2008.",
        "[9] 刘云生. 医院信息系统设计与应用[M]. 北京: 人民卫生出版社, 2018.",
        "[10] 国家卫生健康委员会. 电子病历系统应用水平分级评价标准（试行）[S].",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False)
    page_break(doc)
    add_heading(doc, "致谢", 1)
    add_para(doc, "在本次毕业设计完成过程中，指导教师在选题确定、需求分析、系统设计和论文撰写等方面给予了耐心指导和帮助。通过本次设计，本文作者对 Qt 桌面应用开发、MySQL 数据库设计、客户端服务端通信、业务流程建模和软件测试方法有了更加系统的理解。")
    add_para(doc, "同时，感谢同学和朋友在系统试用、问题反馈和资料整理过程中提供的帮助。由于个人能力和时间有限，系统仍存在不足之处，后续将继续完善业务细节、界面体验和系统稳定性。")


def remove_comments_parts(docx_path: Path):
    # The generated thesis should be clean. Template comments are preserved in the original file only.
    import zipfile
    tmp = docx_path.with_suffix(".tmp.docx")
    comment_patterns = [
        re.compile(r"word/comments.*\.xml$"),
        re.compile(r"word/_rels/comments.*\.xml\.rels$"),
    ]
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if any(p.search(item.filename) for p in comment_patterns):
                continue
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                text = re.sub(r'<Override PartName="/word/comments[^"]+" ContentType="[^"]+"\s*/>', "", text)
                data = text.encode("utf-8")
            if item.filename == "word/_rels/document.xml.rels":
                text = data.decode("utf-8")
                text = re.sub(r'<Relationship [^>]*Target="comments[^"]*"[^>]*/>', "", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"模板不存在：{TEMPLATE}")
    if not STANDARD.exists():
        raise FileNotFoundError(f"撰写规范不存在：{STANDARD}")

    shutil.copyfile(TEMPLATE, WORKING)
    doc = Document(str(WORKING))
    clear_document_body(doc)
    configure_doc(doc)

    add_cover(doc)
    add_abstracts(doc)
    add_toc_placeholder(doc)
    page_break(doc)
    add_chapter_1(doc)
    page_break(doc)
    add_chapter_2(doc)
    page_break(doc)
    add_chapter_3(doc)
    page_break(doc)
    add_chapter_4(doc)
    page_break(doc)
    add_chapter_5(doc)
    page_break(doc)
    add_chapter_6(doc)
    page_break(doc)
    add_references_and_ack(doc)

    doc.save(str(OUTPUT))
    remove_comments_parts(OUTPUT)
    if WORKING.exists():
        WORKING.unlink()
    print(OUTPUT)


if __name__ == "__main__":
    main()
