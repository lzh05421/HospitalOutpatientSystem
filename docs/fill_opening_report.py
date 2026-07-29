from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
INPUT_REPORT = Path("E:/15751/Desktop/开题报告.docx")
OUT_REPORT = Path("E:/15751/Desktop/开题报告-补充完善版.docx")
OUT_REPORT_COPY = ROOT / "开题报告-补充完善版.docx"
OUT_STRUCTURE = ROOT / "系统功能结构图-重绘版.png"

TITLE = "基于 Qt 和 MySQL 的医院门诊挂号与药品管理系统"


def font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def wrap_text(draw, text, font_obj, max_width):
    lines = []
    current = ""
    for ch in text:
        candidate = current + ch
        box = draw.textbbox((0, 0), candidate, font=font_obj)
        if box[2] - box[0] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_center(draw, box, text, font_obj, fill):
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap_text(draw, raw, font_obj, x2 - x1 - 24))
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font_obj)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1] + 6)
    total_h = sum(heights) - 6 if heights else 0
    y = y1 + max(0, (y2 - y1 - total_h) / 2)
    for i, line in enumerate(lines):
        draw.text((x1 + (x2 - x1 - widths[i]) / 2, y), line, font=font_obj, fill=fill)
        y += heights[i]


def draw_structure(path):
    width, height = 2600, 1260
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    root_font = font(42, True)
    module_font = font(27, True)
    leaf_font = font(22)
    caption_font = font(26)
    line = (45, 65, 92)
    border = (55, 76, 106)
    root_fill = (225, 238, 248)
    module_fill = (234, 244, 241)
    leaf_fill = (251, 253, 255)
    head_fill = (212, 231, 226)

    modules = [
        ("登录权限", ["患者预约", "人员登录", "角色权限", "操作日志"]),
        ("院长驾驶舱", ["今日挂号", "待缴账单", "收入概览"]),
        ("患者病历", ["患者建档", "信息维护", "病历档案"]),
        ("科室医生", ["科室目录", "医生信息", "状态管理"]),
        ("排班号源", ["排班维护", "智能补班", "停诊规则"]),
        ("挂号候诊", ["预约挂号", "挂号管理", "候诊叫号"]),
        ("诊疗检查", ["医生接诊", "病历记录", "检查检验"]),
        ("处方药房", ["处方开立", "处方审核", "确认发药"]),
        ("药品库存", ["药品维护", "扫码入库", "库存预警"]),
        ("收费结算", ["账单查询", "缴费退费", "支付记录"]),
        ("统计分析", ["日收入", "科室统计", "图表展示"]),
        ("系统运维", ["配置管理", "多端构建", "MySQL连接"]),
    ]

    root = (740, 46, 1860, 120)
    draw.rounded_rectangle(root, radius=10, fill=root_fill, outline=border, width=3)
    draw_center(draw, root, "医院门诊挂号与药品管理系统", root_font, (16, 35, 58))

    cols = 6
    card_w, card_h = 380, 315
    gap_x, gap_y = 38, 90
    start_x, start_y = 55, 265
    trunk_y = 190
    second_trunk_y = start_y + card_h + 45
    centers = []
    for idx in range(len(modules)):
        row, col = divmod(idx, cols)
        x = start_x + col * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        centers.append((x + card_w // 2, y))

    draw.line((width // 2, root[3], width // 2, second_trunk_y), fill=line, width=3)
    draw.line((centers[0][0], trunk_y, centers[5][0], trunk_y), fill=line, width=3)
    draw.line((centers[6][0], second_trunk_y, centers[11][0], second_trunk_y), fill=line, width=3)

    for idx, (name, leaves) in enumerate(modules):
        row, col = divmod(idx, cols)
        x = start_x + col * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        cx = x + card_w // 2
        branch_y = trunk_y if row == 0 else second_trunk_y
        draw.line((cx, branch_y, cx, y), fill=line, width=3)

        card = (x, y, x + card_w, y + card_h)
        draw.rounded_rectangle(card, radius=8, fill=module_fill, outline=border, width=2)
        head = (x, y, x + card_w, y + 58)
        draw.rounded_rectangle(head, radius=8, fill=head_fill, outline=head_fill, width=0)
        draw_center(draw, head, name, module_font, (16, 66, 70))
        for j, item in enumerate(leaves):
            leaf = (x + 42, y + 86 + j * 54, x + card_w - 42, y + 128 + j * 54)
            draw.rounded_rectangle(leaf, radius=5, fill=leaf_fill, outline=(165, 182, 202), width=1)
            draw_center(draw, leaf, item, leaf_font, (31, 43, 60))

    draw_center(draw, (0, 1166, width, 1216), "图1 系统功能结构图", caption_font, (20, 32, 46))
    image.save(path)


BACKGROUND = [
    ("研究背景与意义：", "heading"),
    ("（一）研究背景", "subheading"),
    ("随着公立医院高质量发展和智慧医院建设持续推进，门诊服务正在从传统窗口排队逐步转向线上线下一体化协同。国务院办公厅在《关于推动公立医院高质量发展的意见》中提出，要推动信息技术与医疗服务深度融合，国家卫生健康委和国家中医药管理局发布的《公立医院高质量发展促进行动（2021—2025年）》也将信息化建设作为提升医疗服务和医院管理能力的重要支撑[1][2]。在医院门诊场景中，患者就诊通常经过患者建档、科室与医生选择、预约挂号、候诊叫号、医生接诊、检查检验、处方开立、药品发放、收费结算和费用统计等环节。各环节数据关联紧密、状态变化频繁，如果仍依赖纸质单据或简单表格管理，容易出现重复录入、号源不同步、库存更新滞后、收费统计不准确和责任追溯困难等问题。", "body"),
    ("从业务需求看，挂号系统不应只停留在保存一条预约记录，还需要与医生排班、时段号源、候诊队列和收费账单保持一致；药品管理也不只是登记库存数量，还应覆盖药品基础信息、分类、扫码入库、库存预警、处方审核与出库记录。门诊业务具有流程短、频次高、窗口并发多、数据变更快等特点，因此系统设计需要兼顾可用性、准确性和可维护性。本课题选择医院门诊挂号与药品管理作为研究对象，能够覆盖医院信息系统中较典型的患者服务、诊疗协同、药品流转和费用管理流程。", "body"),
    ("从技术实现看，Qt Widgets 具备成熟的跨平台桌面界面开发能力，适合构建稳定、响应清晰的业务型桌面应用；Qt SQL 可通过 QODBC 或 QMYSQL 访问 MySQL 数据库，便于将界面操作、业务规则和数据持久化连接起来[5][6]。MySQL 作为成熟的关系型数据库，适合保存患者、医生、排班、挂号、病历、检查、处方、药品、库存、账单、支付、统计和日志等结构化数据[7]。本项目采用 Qt/C++ 实现客户端和服务端，服务端通过 TCP + JSON Lines 协议提供业务接口和权限校验，数据库负责持久化保存业务数据，能够体现客户端、服务端、数据库和工程构建的综合实践能力。", "body"),
    ("（二）研究意义", "subheading"),
    ("本课题的意义主要体现在三个方面。第一，在业务应用层面，系统围绕门诊主流程设计，将患者信息、医生排班、挂号号源、候诊叫号、接诊病历、检查检验、处方药品、库存变化、收费结算和统计分析统一管理，有助于减少信息孤岛和重复录入，提高门诊窗口处理效率。第二，在数据库设计层面，系统通过角色、用户、科室、医生、患者、排班、挂号、病历、检查、处方、药品、库存、账单、支付、统计和审计日志等表建立主外键关系，能够训练较完整的关系数据库建模、查询、事务和数据一致性设计能力。第三，在工程实践层面，本课题需要完成 Qt 界面设计、网络通信、服务端路由、角色权限控制、MySQL 访问、分页查询、自动刷新、统计图表、异常处理和跨平台构建，能够完整体现软件工程中的需求分析、概要设计、详细设计、编码实现、测试验证和部署维护过程。与单一模块的小型管理程序相比，本系统功能链条更完整，更接近真实医院门诊信息系统的基础形态，具有较好的实践价值和教学价值。", "body"),
]


MAIN_CONTENT = [
    ("主要内容：", "heading"),
    ("本系统面向医院门诊挂号与药品管理业务，采用客户端/服务器端架构进行设计。客户端负责患者预约入口和医院人员业务界面，服务器端负责统一接口、权限校验和业务处理，MySQL 数据库负责保存核心业务数据。结合当前项目代码和业务页面，系统主要分为十二个模块、三十七个子功能，系统功能结构图如图1所示。", "body"),
    ("__PICTURE__", "picture"),
    ("（一）登录权限管理模块", "subheading"),
    ("该模块包括患者预约、医院人员登录、角色权限和操作日志四个子功能。患者预约入口面向普通就诊人员，医院人员登录入口面向管理员、科主任、挂号员、医生、药房人员和收费员；服务端根据角色控制可访问菜单与接口，并对新增、修改、删除、叫号、审核、收费等关键操作进行日志记录。", "body"),
    ("（二）院长驾驶舱模块", "subheading"),
    ("该模块包括今日挂号、待缴账单和收入概览三个子功能。系统通过汇总挂号、收费和库存预警等数据，为管理员、科主任和收费人员提供门诊运行概览，便于快速了解当天业务状态。", "body"),
    ("（三）患者病历管理模块", "subheading"),
    ("该模块包括患者建档、信息维护和病历档案三个子功能。患者建档维护患者编号、姓名、性别、身份证号、联系电话和地址等资料；病历档案关联挂号、接诊、检查、处方和费用记录，便于医生回看患者历史就诊情况。", "body"),
    ("（四）科室医生管理模块", "subheading"),
    ("该模块包括科室目录、医生信息和状态管理三个子功能。科室管理维护门诊大类、专科和诊室信息，医生管理维护医生所属科室、职称、专长、挂号费和在职状态，为排班、挂号和统计提供基础数据。", "body"),
    ("（五）排班号源管理模块", "subheading"),
    ("该模块包括排班维护、智能补班和停诊规则三个子功能。管理员或科主任可按医生、日期和时段维护出诊安排，设置总号源、剩余号源和出诊状态；挂号成功后系统扣减对应号源，医生停诊或临时调整时可及时更新排班。", "body"),
    ("（六）挂号候诊管理模块", "subheading"),
    ("该模块包括预约挂号、挂号管理和候诊叫号三个子功能。患者或挂号员可根据科室、医生、日期和时段完成挂号，系统生成挂号记录并维护等待、已叫号、检查后候诊、已完成、已取消等状态；候诊队列支持按科室、医生和状态筛选。", "body"),
    ("（七）诊疗检查管理模块", "subheading"),
    ("该模块包括医生接诊、病历记录和检查检验三个子功能。医生可查看候诊患者，填写主诉、现病史、既往史、体征、诊断和医嘱；需要检查时可开立检查申请，检查结果回传后患者可继续进入接诊流程。", "body"),
    ("（八）处方药房管理模块", "subheading"),
    ("该模块包括处方开立、处方审核和确认发药三个子功能。医生根据诊断结果开立处方并添加药品明细，药房人员审核处方后确认发药，系统将处方金额、药品库存和收费账单联动处理。", "body"),
    ("（九）药品库存管理模块", "subheading"),
    ("该模块包括药品维护、扫码入库和库存预警三个子功能。系统维护药品编码、条形码、名称、分类、规格、单位、进价、售价、库存数量、预警数量和有效期，支持扫码或手动入库，并根据库存阈值提示补药。", "body"),
    ("（十）收费结算管理模块", "subheading"),
    ("该模块包括账单查询、缴费退费和支付记录三个子功能。系统根据挂号费、药品费和检查费生成账单，支持未缴费、已缴费、已退费等状态管理，并记录支付方式、支付金额、收费员和收费时间。", "body"),
    ("（十一）统计分析模块", "subheading"),
    ("该模块包括日收入、科室统计和图表展示三个子功能。系统按日期、科室和费用类型统计挂号收入、药品收入和总收入，通过表格和图表展示门诊费用构成，为管理人员分析业务情况提供依据。", "body"),
    ("（十二）系统运维模块", "subheading"),
    ("该模块包括配置管理、多端构建和 MySQL 连接三个子功能。项目使用 CMake 组织 client、server、common 和 launcher 等子工程，支持 VS Code、Qt Creator、Windows、Linux 和银河麒麟环境构建运行；服务端通过配置文件选择数据库连接参数，也保留演示模式，便于课堂演示和环境异常时验证系统功能。", "body"),
]


PLAN = [
    ("工作方案及进度安排：", "heading"),
    ("（一）工作方案", "subheading"),
    ("本课题按照软件工程开发流程推进。首先进行需求分析，梳理患者预约、医院人员登录、医生排班、挂号候诊、医生接诊、检查检验、处方药房、药品库存、收费结算、统计分析和操作日志等功能边界；其次进行数据库设计，建立角色、用户、科室、医生、患者、排班、挂号、病历、检查、处方、药品、库存、账单、支付、统计和审计日志等数据表，并确定主外键关系与核心业务编号；然后进行系统架构设计，采用 Qt Widgets 实现客户端界面，采用 Qt TCP 服务端提供接口，使用 JSON Lines 作为网络消息格式，使用 MySQL 进行数据持久化；最后进行编码实现、联调测试、部署验证和论文撰写。", "body"),
    ("技术路线为：Qt Widgets 客户端界面 → ApiClient 封装 TCP 请求 → Protocol 编解码 JSON Lines → HospitalServer 接收连接 → RequestRouter 进行权限校验和服务分发 → 各业务 Service 处理请求 → DatabaseManager 连接 MySQL → 返回统一响应。测试阶段将按照患者、挂号员、医生、药房人员、收费员、科主任和管理员等角色分别验证，重点关注号源扣减是否正确、挂号新增后医院端是否刷新显示、处方与药品库存是否联动、账单支付状态是否一致、统计数据是否可追溯、断开数据库后的错误提示是否清晰。", "body"),
    ("（二）进度安排", "subheading"),
    ("2025年12月08日---2025年12月21日：完成课题调研、需求分析、系统功能模块划分和开题报告撰写。", "body"),
    ("2025年12月22日---2026年01月04日：完成系统总体设计，包括 C/S 架构、数据库表结构、主要业务流程和界面原型。", "body"),
    ("2026年01月05日---2026年02月08日：完成项目基础框架，实现客户端登录入口、服务端通信框架、MySQL 连接配置和基础数据初始化。", "body"),
    ("2026年02月09日---2026年03月01日：完成患者管理、挂号预约、挂号管理、候诊队列和医生排班等核心模块，实现号源限制、刷新和分页查询。", "body"),
    ("2026年03月02日---2026年03月22日：完成科室医生管理、医生接诊、病历档案、检查检验和处方管理模块，实现诊疗业务闭环。", "body"),
    ("2026年03月23日---2026年04月12日：完成药品库存、扫码入库、处方审核发药、收费结算和费用统计模块，实现库存、账单和统计联动。", "body"),
    ("2026年04月13日---2026年04月26日：进行系统联调和功能测试，修复权限控制、数据库连接、页面刷新、分页显示和异常提示问题。", "body"),
    ("2026年04月27日---2026年05月10日：完成 Windows、Linux、银河麒麟、Qt Creator 和 VS Code 环境下的构建运行验证，整理部署说明。", "body"),
    ("2026年05月11日---2026年05月24日：撰写毕业设计论文，完善系统截图、数据库设计说明、核心代码说明和测试分析。", "body"),
    ("2026年05月25日---2026年06月05日：根据指导教师意见修改论文和系统，准备答辩 PPT，完成最终提交。", "body"),
]


REFERENCES = [
    "[1] 国务院办公厅. 关于推动公立医院高质量发展的意见[Z]. 2021.",
    "[2] 国家卫生健康委, 国家中医药管理局. 公立医院高质量发展促进行动（2021—2025年）[Z]. 2021.",
    "[3] 国家卫生健康委, 国家中医药管理局, 国家疾病预防控制局. 医疗卫生机构网络安全管理办法[Z]. 2022.",
    "[4] 国家卫生健康委. 三级医院评审标准（2022年版）及其实施细则[S]. 2022.",
    "[5] The Qt Company. Qt Widgets Documentation[EB/OL]. https://doc.qt.io/qt-6/qtwidgets-index.html, 2026-06-08.",
    "[6] The Qt Company. QSqlDatabase Class Documentation[EB/OL]. https://doc.qt.io/qt-6/qsqldatabase.html, 2026-06-08.",
    "[7] Oracle Corporation. MySQL 8.4 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.4/en/, 2026-06-08.",
    "[8] Kitware. CMake Reference Documentation[EB/OL]. https://cmake.org/cmake/help/latest/, 2026-06-08.",
    "[9] HL7 International. FHIR Release 5 Specification[EB/OL]. https://hl7.org/fhir/, 2026-06-08.",
    "[10] 王珊, 萨师煊. 数据库系统概论[M]. 第6版. 北京: 高等教育出版社, 2023.",
    "[11] 张海藩, 牟永敏. 软件工程导论[M]. 第7版. 北京: 清华大学出版社, 2022.",
    "[12] Dey N. Cross-Platform Development with Qt 6 and Modern C++[M]. Birmingham: Packt Publishing, 2021.",
]


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def clear_cell(cell):
    cell.text = ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""


def add_formatted_paragraph(cell, text, kind="body"):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if kind == "heading":
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, 12, True)
    elif kind == "subheading":
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, 12, True)
    else:
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run(text)
        set_run_font(run, 12, False)
    return p


def add_picture(cell, image_path):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.8))

    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.first_line_indent = Pt(0)
    run = cap.add_run("图1 系统功能结构图")
    set_run_font(run, 10.5, False)


def fill_content_cell(cell, blocks):
    clear_cell(cell)
    for text, kind in blocks:
        if kind == "picture":
            add_picture(cell, OUT_STRUCTURE)
        else:
            add_formatted_paragraph(cell, text, kind)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def fill_references(cell):
    clear_cell(cell)
    add_formatted_paragraph(cell, "四、参考文献：", "heading")
    for item in REFERENCES:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.line_spacing = Pt(23)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, 12, False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_short_cell(cell, text=None):
    if text is not None:
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        set_run_font(run, 12, False)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, 12, run.bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{name}"))
        if elem is None:
            elem = OxmlElement(f"w:{name}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def build_report():
    if not INPUT_REPORT.exists():
        raise FileNotFoundError(f"未找到原开题报告：{INPUT_REPORT}")

    draw_structure(OUT_STRUCTURE)

    doc = Document(str(INPUT_REPORT))
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, 12, run.bold)

    table = doc.tables[0]
    set_table_borders(table)
    for row_idx in range(4):
        for cell in table.rows[row_idx].cells:
            set_short_cell(cell)
    set_short_cell(table.cell(3, 1), TITLE)

    fill_content_cell(table.cell(4, 0), BACKGROUND)
    fill_content_cell(table.cell(5, 0), MAIN_CONTENT)
    fill_content_cell(table.cell(6, 0), PLAN)
    fill_references(table.cell(7, 0))

    doc.save(str(OUT_REPORT))
    shutil.copy2(OUT_REPORT, OUT_REPORT_COPY)
    print(OUT_REPORT)
    print(OUT_STRUCTURE)


if __name__ == "__main__":
    build_report()
