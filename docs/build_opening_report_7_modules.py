from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
INPUT_REPORT = Path("E:/15751/Desktop/开题报告.docx")
OUT_REPORT = Path("E:/15751/Desktop/开题报告-七模块新版.docx")
OUT_REPORT_COPY = ROOT / "开题报告-七模块新版.docx"
OUT_STRUCTURE = ROOT / "系统功能结构图-七模块新版.png"

TITLE = "基于 Qt 和 MySQL 的医院门诊挂号与药品管理系统的设计与实现"


def font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf") if bold else Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def wrap_text(draw, text, font_obj, max_width):
    lines = []
    current = ""
    for ch in text:
        candidate = current + ch
        box = draw.textbbox((0, 0), candidate, font=font_obj)
        if box[2] - box[0] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_center(draw, box, text, font_obj, fill):
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap_text(draw, raw, font_obj, x2 - x1 - 28))
    metrics = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font_obj)
        metrics.append((line, bbox[2] - bbox[0], bbox[3] - bbox[1] + 8))
    total_h = sum(item[2] for item in metrics) - 8 if metrics else 0
    y = y1 + max(0, (y2 - y1 - total_h) / 2)
    for line, width, height in metrics:
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font_obj, fill=fill)
        y += height


def draw_structure(path):
    width, height = 2600, 1220
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    root_font = font(42, True)
    module_font = font(28, True)
    leaf_font = font(22)
    caption_font = font(26)
    line = (48, 66, 94)
    border = (55, 76, 106)
    root_fill = (226, 238, 249)
    module_fill = (236, 245, 242)
    head_fill = (214, 232, 228)
    leaf_fill = (251, 253, 255)

    root = (700, 45, 1900, 122)
    draw.rounded_rectangle(root, radius=10, fill=root_fill, outline=border, width=3)
    draw_center(draw, root, "医院门诊挂号与药品管理系统", root_font, (15, 35, 58))

    modules = [
        ("用户登录与权限管理", ["患者注册登录", "医院人员登录", "角色权限控制", "权限配置", "操作日志"]),
        ("患者与就诊人管理", ["患者建档", "就诊人管理", "信息维护", "病历档案", "历史订单"]),
        ("门诊预约挂号", ["智能分诊", "科室医生选择", "号源查询", "预约/急诊挂号", "医保预校验"]),
        ("医生排班与候诊接诊", ["排班维护", "号源规则", "候诊队列", "叫号接诊", "结构化病历"]),
        ("检查检验与处方管理", ["检查申请", "结果录入", "处方开立", "审核/驳回", "发药/退药"]),
        ("药品库存管理", ["药品信息", "药品分类", "扫码/条码入库", "库存预警", "PASS规则"]),
        ("收费结算与统计分析", ["账单生成", "自费/医保支付", "退款审核", "支付记录", "驾驶舱统计"]),
    ]

    cols = 4
    card_w, card_h = 575, 310
    gap_x, gap_y = 48, 100
    start_x, start_y = 80, 270
    trunk_y = 185
    row_trunk_y = start_y + card_h + 54
    draw.line((width // 2, root[3], width // 2, row_trunk_y), fill=line, width=4)
    draw.line((start_x + card_w // 2, trunk_y, start_x + 3 * (card_w + gap_x) + card_w // 2, trunk_y), fill=line, width=4)
    draw.line((start_x + card_w // 2, row_trunk_y, start_x + 2 * (card_w + gap_x) + card_w // 2, row_trunk_y), fill=line, width=4)

    for idx, (name, leaves) in enumerate(modules):
        row, col = divmod(idx, cols)
        x = start_x + col * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        cx = x + card_w // 2
        branch_y = trunk_y if row == 0 else row_trunk_y
        draw.line((cx, branch_y, cx, y), fill=line, width=4)
        card = (x, y, x + card_w, y + card_h)
        draw.rounded_rectangle(card, radius=8, fill=module_fill, outline=border, width=3)
        head = (x, y, x + card_w, y + 66)
        draw.rounded_rectangle(head, radius=8, fill=head_fill, outline=head_fill, width=0)
        draw_center(draw, head, name, module_font, (18, 64, 66))
        for j, item in enumerate(leaves):
            leaf = (x + 44, y + 90 + j * 42, x + card_w - 44, y + 125 + j * 42)
            draw.rounded_rectangle(leaf, radius=5, fill=leaf_fill, outline=(168, 183, 202), width=1)
            draw_center(draw, leaf, item, leaf_font, (31, 43, 60))

    draw_center(draw, (0, 1132, width, 1185), "图1 系统功能结构图", caption_font, (20, 32, 46))
    image.save(path)


BACKGROUND = [
    ("研究背景与意义：", "heading"),
    ("（一）研究背景", "subheading"),
    ("随着公立医院高质量发展和智慧医院建设持续推进，门诊服务正在从传统窗口排队逐步转向线上线下一体化协同。国务院办公厅在《关于推动公立医院高质量发展的意见》中提出，要推动信息技术与医疗服务深度融合；国家卫生健康委等部门发布的《公立医院高质量发展促进行动（2021—2025年）》也将信息化建设作为提升医疗服务和医院管理能力的重要支撑[1][2]。在医院门诊场景中，患者就诊通常需要经过患者建档、科室与医生选择、预约挂号、候诊叫号、医生接诊、检查检验、处方开立、药品发放、收费结算和费用统计等环节。各环节数据关联紧密、状态变化频繁，如果仍依赖纸质单据或简单表格管理，容易出现重复录入、号源不同步、库存更新滞后、收费统计不准确和责任追溯困难等问题。", "body"),
    ("从业务需求看，挂号系统不应只停留在保存预约记录，还需要与医生排班、时段号源、候诊队列和收费账单保持一致；药品管理也不只是登记库存数量，还应覆盖药品基础信息、分类维护、扫码入库、库存预警、处方审核与药品出库记录。门诊业务具有流程短、频次高、窗口并发多、数据变更快等特点，因此系统设计需要兼顾可用性、准确性和可维护性。本课题选择医院门诊挂号与药品管理作为研究对象，能够覆盖医院信息系统中较典型的患者服务、诊疗协同、药品流转和费用管理流程。", "body"),
    ("从技术实现看，Qt 具有成熟的跨平台桌面界面开发能力，Qt Widgets 适合构建稳定、响应清晰的业务型桌面应用；Qt SQL 可通过 QODBC 或 QMYSQL 访问 MySQL 数据库，便于将界面操作、业务规则和数据持久化连接起来[5][6]。MySQL 作为成熟的关系型数据库，适合保存患者、医生、排班、挂号、病历、检查、处方、药品、库存、账单、支付、统计和日志等结构化数据[7]。本项目采用 Qt/C++ 实现客户端和服务端，服务端通过 TCP 与 JSON Lines 协议提供业务接口和权限校验，数据库负责持久化保存业务数据，能够体现客户端、服务端、数据库和工程构建的综合实践能力。", "body"),
    ("（二）研究意义", "subheading"),
    ("本课题的研究意义主要体现在三个方面。第一，在业务应用层面，系统围绕门诊主流程设计，将患者信息、医生排班、挂号号源、候诊叫号、接诊病历、检查检验、处方药品、库存变化、收费结算和统计分析统一管理，有助于减少信息孤岛和重复录入，提高门诊窗口处理效率。第二，在数据库设计层面，系统通过角色、用户、科室、医生、患者、排班、挂号、病历、检查、处方、药品、库存、账单、支付、统计和审计日志等表建立主外键关系，能够训练较完整的关系数据库建模、查询、事务和数据一致性设计能力。第三，在工程实践层面，本课题需要完成 Qt 界面设计、网络通信、服务端路由、角色权限控制、MySQL 访问、分页查询、自动刷新、统计图表、异常处理和跨平台构建，能够完整体现软件工程中的需求分析、概要设计、详细设计、编码实现、测试验证和部署维护过程。与单一模块的小型管理程序相比，本系统功能链条更完整，更接近真实医院门诊信息系统的基础形态，具有较好的实践价值和教学价值。", "body"),
]


MAIN_CONTENT = [
    ("主要内容：", "heading"),
    ("本系统面向医院门诊挂号与药品管理业务，采用客户端/服务端结构进行设计。客户端负责患者预约入口和医院人员业务界面，服务端负责统一接口、权限校验和业务处理，MySQL 数据库负责保存核心业务数据。为避免功能表达过于分散，开题报告按照业务流程将项目归纳为七个主模块、三十五个子功能，系统功能结构图如图1所示。", "body"),
    ("__PICTURE__", "picture"),
    ("（一）用户登录与权限管理模块", "subheading"),
    ("该模块包括患者注册登录、医院人员登录、角色权限控制、权限配置和操作日志五个子功能。患者入口面向普通就诊人员，支持患者账号登录、注册和个人预约；医院人员入口面向管理员、科主任、挂号员、医生、药房人员和收费员等角色。系统根据角色控制可访问菜单和可执行操作，并记录新增、修改、删除、叫号、审核、收费等关键操作日志，便于系统使用过程追溯。", "body"),
    ("（二）患者与就诊人管理模块", "subheading"),
    ("该模块包括患者建档、就诊人管理、患者信息维护、病历档案和历史订单五个子功能。医院人员可维护患者姓名、性别、出生日期、身份证号、联系电话、家庭地址等基础信息；患者端可管理本人或家属就诊人信息，并查看历史预约订单。患者数据与挂号记录、接诊病历、检查结果、处方明细和收费账单关联，是系统业务流转的基础。", "body"),
    ("（三）门诊预约挂号模块", "subheading"),
    ("该模块包括智能分诊、科室医生选择、号源查询、预约/急诊挂号和医保资格预校验五个子功能。患者可根据症状描述进行科室推荐，也可手动选择门诊大类、专科、诊室、医生、日期和时段。系统根据医生排班与剩余号源展示可预约信息，挂号成功后生成挂号记录并扣减对应号源；急诊挂号记录急诊原因，医保统筹挂号前进行资格预校验，降低后续收费环节的状态冲突。", "body"),
    ("（四）医生排班与候诊接诊模块", "subheading"),
    ("该模块包括排班维护、号源规则、候诊队列、叫号接诊和结构化病历五个子功能。管理员或科主任可维护医生出诊日期、时段、总号源和剩余号源，并设置停诊或排班规则。挂号成功的患者进入候诊队列，医生或挂号员可进行叫号，医生接诊时填写主诉、现病史、既往史、体征、诊断和医嘱等结构化病历内容。", "body"),
    ("（五）检查检验与处方管理模块", "subheading"),
    ("该模块包括检查申请、检查结果录入、处方开立、处方审核/驳回和发药/退药五个子功能。医生可在接诊过程中开立检查检验申请，检查完成后录入结果；根据诊断结果开具处方并填写药品数量、用法用量、频次和天数。药房人员可对处方进行审核、驳回、发药或退药，处方状态与收费和库存业务保持联动。", "body"),
    ("（六）药品库存管理模块", "subheading"),
    ("该模块包括药品信息、药品分类、扫码/条码入库、库存预警和 PASS 用药规则五个子功能。系统维护药品编码、条形码、药品名称、分类、规格、单位、进价、售价、库存数量、预警数量和有效期等信息。药品入库既支持手动录入，也支持条码识别；库存低于预警数量时进行提示。PASS 用药规则用于对过敏、剂量和联合用药等风险进行提醒或阻断，提高处方审核的规范性。", "body"),
    ("（七）收费结算与统计分析模块", "subheading"),
    ("该模块包括账单生成、自费/医保支付、退款审核、支付记录和驾驶舱统计五个子功能。系统根据挂号费、检查费和药品费生成账单，支持自费支付、医保支付、模拟扫码支付和支付状态查询；退费业务通过申请和审核流程进行控制。院长驾驶舱和费用统计页面汇总今日挂号、待缴账单、药品库存预警、科室收入和费用构成等数据，为医院管理人员掌握门诊运行情况提供参考。", "body"),
]


PLAN = [
    ("工作方案及进度安排：", "heading"),
    ("（一）工作方案", "subheading"),
    ("本课题按照软件工程开发流程推进。首先进行需求分析，梳理患者预约、医院人员登录、医生排班、挂号候诊、医生接诊、检查检验、处方药房、药品库存、收费结算、统计分析和操作日志等功能边界；其次进行数据库设计，建立角色、用户、科室、医生、患者、排班、挂号、病历、检查、处方、药品、库存、账单、支付、统计和审计日志等数据表，并确定主外键关系与核心业务编号；然后进行系统架构设计，采用 Qt Widgets 实现客户端界面，采用 Qt TCP 服务端提供接口，使用 JSON Lines 作为网络消息格式，使用 MySQL 进行数据持久化；最后进行编码实现、联调测试、部署验证和论文撰写。", "body"),
    ("技术路线为：Qt Widgets 客户端界面、ApiClient 封装 TCP 请求、Protocol 编解码 JSON Lines、HospitalServer 接收连接、RequestRouter 进行权限校验和服务分发、各业务 Service 处理请求、DatabaseManager 连接 MySQL、返回统一响应。测试阶段将按照患者、挂号员、医生、药房人员、收费员、科主任和管理员等角色分别验证，重点关注号源扣减是否正确、挂号新增后医院端是否刷新显示、处方与药品库存是否联动、账单支付状态是否一致、统计数据是否可追溯、断开数据库后的错误提示是否清晰。", "body"),
    ("（二）进度安排", "subheading"),
    ("2026年06月14日——2026年06月28日：系统需求分析、确定功能模块、完成开题报告。", "body"),
    ("2026年06月29日——2026年07月05日：完成开题答辩，根据教师意见调整题目、模块范围和功能结构图。", "body"),
    ("2026年07月06日——2026年07月11日：搭建项目框架，完成客户端、服务端、公共协议和构建脚本的概要设计。", "body"),
    ("2026年07月12日——2026年07月22日：完成数据库概念及逻辑设计，明确患者、医生、排班、挂号、病历、处方、药品、账单和日志等表关系。", "body"),
    ("2026年07月23日——2026年08月08日：完成系统主要功能的实现，包括患者管理、预约挂号、排班号源、候诊接诊、处方药品和收费结算。", "body"),
    ("2026年08月09日——2026年08月29日：完善项目功能及撰写毕业论文，补充医保支付、检查检验、退款审核、统计图表和权限日志等功能。", "body"),
    ("2026年08月30日——2026年09月06日：进行系统测试，完成中期报告，重点验证多角色操作、数据库连接、号源扣减、库存更新和费用统计。", "body"),
    ("2026年09月07日——2026年09月13日：中期检查、答辩，根据反馈修正系统问题和论文结构。", "body"),
    ("2026年09月14日——2026年09月24日：完善项目功能与毕业论文，补齐系统截图、数据库设计图和关键流程说明。", "body"),
    ("2026年09月25日——2026年09月27日：论文初稿定稿，保证系统主要流程能够稳定运行。", "body"),
    ("2026年09月28日——2026年10月04日：答辩稿定稿、重复率检测、AIGC检测。", "body"),
    ("2026年10月05日——2026年10月11日：指导教师评阅、评阅教师评阅，根据意见进行修改。", "body"),
    ("2026年10月12日——2026年10月18日：模拟答辩、终稿、检测，整理最终提交材料。", "body"),
    ("2027年05月24日——2027年05月30日：完成毕业答辩。", "body"),
]


REFERENCES = [
    "[1] 国务院办公厅. 关于推动公立医院高质量发展的意见[Z]. 2021.",
    "[2] 国家卫生健康委, 国家中医药管理局. 公立医院高质量发展促进行动（2021—2025年）[Z]. 2021.",
    "[3] 国家卫生健康委, 国家中医药管理局, 国家疾病预防控制局. 医疗卫生机构网络安全管理办法[Z]. 2022.",
    "[4] 国家卫生健康委. 三级医院评审标准（2022年版）及其实施细则[S]. 2022.",
    "[5] The Qt Company. Qt Widgets Documentation[EB/OL]. https://doc.qt.io/qt-6/qtwidgets-index.html, 2026.",
    "[6] The Qt Company. QSqlDatabase Class Documentation[EB/OL]. https://doc.qt.io/qt-6/qsqldatabase.html, 2026.",
    "[7] Oracle Corporation. MySQL 8.4 Reference Manual[EB/OL]. https://docs.oracle.com/cd/E17952_01/mysql-8.4-en/, 2026.",
    "[8] Kitware. CMake Documentation[EB/OL]. https://cmake.org/cmake/help/latest/, 2026.",
    "[9] HL7 International. FHIR Release 5 Specification[EB/OL]. https://hl7.org/fhir/R5/, 2023.",
    "[10] 王珊, 萨师煊. 数据库系统概论[M]. 第6版. 北京: 高等教育出版社, 2023.",
    "[11] 张海藩, 牟永敏. 软件工程导论[M]. 第7版. 北京: 清华大学出版社, 2022.",
    "[12] Dey N. Cross-Platform Development with Qt 6 and Modern C++[M]. Birmingham: Packt Publishing, 2021.",
]


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def clear_cell(cell):
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def add_paragraph(cell, text, kind="body"):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if kind in ("heading", "subheading"):
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, 12, True)
    else:
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run(text)
        set_run_font(run, 12, False)


def add_picture(cell):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(OUT_STRUCTURE), width=Cm(15.8))

    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.first_line_indent = Pt(0)
    run = cap.add_run("图1 系统功能结构图")
    set_run_font(run, 10.5, False)


def fill_content_cell(cell, blocks):
    clear_cell(cell)
    for text, kind in blocks:
        if kind == "picture":
            add_picture(cell)
        else:
            add_paragraph(cell, text, kind)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def fill_references(cell):
    clear_cell(cell)
    add_paragraph(cell, "四、参考文献：", "heading")
    for item in REFERENCES:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.line_spacing = Pt(23)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, 12, False)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_short_cell(cell, text=None):
    if text is not None:
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        set_run_font(run, 12, False)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, 12, run.bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{name}"))
        if elem is None:
            elem = OxmlElement(f"w:{name}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def merged_cell(table, row_idx):
    try:
        return table.cell(row_idx, 0).merge(table.cell(row_idx, 3))
    except Exception:
        return table.cell(row_idx, 0)


def build_report():
    if not INPUT_REPORT.exists():
        raise FileNotFoundError(f"未找到原开题报告：{INPUT_REPORT}")

    draw_structure(OUT_STRUCTURE)

    doc = Document(str(INPUT_REPORT))
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, 12, run.bold)

    table = doc.tables[0]
    set_table_borders(table)
    for row_idx in range(4):
        for cell in table.rows[row_idx].cells:
            set_short_cell(cell)
    set_short_cell(table.cell(3, 1), TITLE)

    fill_content_cell(merged_cell(table, 4), BACKGROUND)
    fill_content_cell(merged_cell(table, 5), MAIN_CONTENT)
    fill_content_cell(merged_cell(table, 6), PLAN)
    fill_references(merged_cell(table, 7))

    guide_cell = merged_cell(table, 8)
    clear_cell(guide_cell)
    add_paragraph(guide_cell, "指导教师意见：", "heading")
    for _ in range(5):
        add_paragraph(guide_cell, "", "body")
    add_paragraph(guide_cell, "指导教师：", "body")
    add_paragraph(guide_cell, "年    月    日", "body")

    major_cell = merged_cell(table, 9)
    clear_cell(major_cell)
    add_paragraph(major_cell, "所在专业意见：", "heading")
    add_paragraph(major_cell, "□通过    □不通过", "body")
    for _ in range(3):
        add_paragraph(major_cell, "", "body")
    add_paragraph(major_cell, "负责人：", "body")
    add_paragraph(major_cell, "年    月    日", "body")

    doc.save(str(OUT_REPORT))
    shutil.copy2(OUT_REPORT, OUT_REPORT_COPY)
    print(OUT_REPORT)
    print(OUT_REPORT_COPY)
    print(OUT_STRUCTURE)


if __name__ == "__main__":
    build_report()
